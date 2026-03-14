#!/usr/bin/env python3
"""
Generate QIP Incentive Policy Document Version 3 (English)
- Comprehensive employee-friendly guide with detailed examples
- Screenshots embedded from policy_screenshots/ directory
- 100% aligned with V10 system code
- Validity: March 2026 ~ February 2027
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Path Configuration ──
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
SCREENSHOT_DIR = os.path.join(PROJECT_DIR, 'policy_screenshots')


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_header_row(table, headers, bg="2F5496"):
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], bg)


def add_row(table, cells_data, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold:
            run.bold = True
        if bg_color:
            set_cell_shading(cell, bg_color)
    return row


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def bold_para(doc, bold_text, normal_text=""):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_screenshot(doc, filename, caption=None, width=Inches(6)):
    """Add a screenshot image with optional caption."""
    img_path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(89, 89, 89)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot not found: {filename}]")
        run.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)


def add_example_box(doc, title, lines):
    """Add a visually distinct example block using a single-cell table with shading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "E8F0FE")
    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(47, 84, 150)
    # Content lines
    for line in lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
    doc.add_paragraph()  # spacing


def create_policy():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ══════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HWK QIP INCENTIVE POLICY")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 3 — Updated March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Validity Period: March 2026 ~ February 2027")
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Quality Improvement Program Department\nHWK Vina")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════
    doc.add_heading("TABLE OF CONTENTS", level=1)

    toc_items = [
        ("Section 1.", " Overview"),
        ("Section 2.", " General Explanation on Incentive Policy Structure"),
        ("Section 3.", " Scope and Objects"),
        ("Section 4.", " Definitions & Abbreviations"),
        ("Section 5.", " Incentive Amount Calculation Method — Program 1"),
        ("Section 6.", " Incentive Amount Calculation Method — Program 2"),
        ("Section 7.", " Summary of Exceptional Clauses"),
        ("Section 8.", " System Architecture & Dashboard (V10)"),
        ("Section 9.", " Quick Reference & FAQ"),
        ("Appendix 1:", " 10 Incentive Conditions — Detailed Guide with Examples"),
        ("Appendix 2:", " Position-Condition Application Matrix"),
        ("Appendix 3:", " HWK QIP Talent Pool Status"),
        ("Appendix 4:", " QIP Team R&R Map"),
        ("Appendix 5:", " Firestore Data Collections"),
    ]
    for bold_part, normal_part in toc_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(normal_part)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 1. OVERVIEW
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 1. OVERVIEW", level=1)

    doc.add_paragraph(
        '"HWK QIP INCENTIVE POLICY" is designed to renew every 6 months. The 1st version was created '
        'in February 2025, and this document represents the updated Version 3, reflecting the current '
        'V10 system architecture with Firestore-based data management and automated calculation pipeline.'
    )
    doc.add_paragraph(
        'This document provides an overview of the HWK company\'s incentive policy aimed at achieving '
        'key quality targets. By offering incentives based on quality performance metrics, we seek to '
        'strengthen company-wide quality management and ultimately reach our 5Q KPI milestones as a '
        'Global Top 3 quality Factory continuously. Through this incentive scheme, we aim to solidify '
        'our HWK factory culture of quality management, increase employee engagement, and reduce turnover '
        'rates, thereby ensuring a stable and motivated workforce. The QIP incentive policy applies to '
        'all QIP team members.'
    )

    doc.add_heading("PURPOSE", level=2)
    bullet(doc, "To improve & achieve the best HWK quality performance through sustainable motivation.")
    bullet(doc, "Establish a continuous quality management system: By providing quality performance-based incentives, "
           "our goal is to maintain and enhance company-wide quality initiatives.")
    bullet(doc, "Achieve 5Q goals: We aim to motivate all teams, including the quality team, to meet the highest "
           "quality standards (5Q) and ensure that the company's objectives are clearly shared across the organization.")
    bullet(doc, "Enhance motivation and organizational stability: Through a well-defined rewards system tied to "
           "excellent quality performance, we seek to increase employees' sense of accomplishment and effectively manage turnover rates.")
    bullet(doc, "Strengthen collaboration and internal communication: By fostering interdepartmental collaboration and "
           "sharing of best practices, we aim to support overall organizational performance and progress toward quality targets.")

    doc.add_heading("HOW TO READ THIS DOCUMENT", level=2)
    doc.add_paragraph(
        "If you are a new employee, we recommend reading this document in the following order:"
    )
    bullet(doc, "First, read Section 2 to understand the overall structure (3 types, 2 programs).")
    bullet(doc, "Then, read Section 3 to find out which type YOU belong to (Type-1, Type-2, or Type-3).")
    bullet(doc, "Next, read Appendix 1 carefully — it explains each of the 10 conditions in detail with real examples.")
    bullet(doc, "Finally, read Section 5 to understand how your incentive amount is calculated.")
    bullet(doc, "Section 8 shows you how to use the online dashboard to check your status.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 2. GENERAL EXPLANATION
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 2. GENERAL EXPLANATION ON INCENTIVE POLICY STRUCTURE", level=1)

    doc.add_paragraph("HWK QIP incentive policy has mainly 2 Programs:")
    bullet(doc, 'Program 1: "Monthly-Performance-Based-Incentive Program"')
    bullet(doc, 'Program 2: "6-Month-Performance-Based-Incentive Program"')

    doc.add_paragraph(
        "Program 1 of HWK QIP incentive policy has 3 evaluation processes. Every QIP member is categorized "
        "into three types: Type-1, Type-2, and Type-3. Depending on a member's type, the incentive calculation "
        "method is applied differently."
    )

    doc.add_heading("Type-1 Evaluation (Direct Assembly QIP)", level=2)
    doc.add_paragraph(
        "For Type-1 QIP team members, performance is evaluated based on up to 10 conditions across four categories:"
    )
    bullet(doc, "Attendance conditions (C1-C4): Attendance Rate, Unapproved Absence, Actual Working Days, Minimum Working Days")
    bullet(doc, "AQL conditions (C5-C8): Personal AQL failure, personal consecutive failures, team/area consecutive failures, area reject rate")
    bullet(doc, "5PRS conditions (C9-C10): 5PRS pass rate, 5PRS inspection quantity")

    doc.add_paragraph(
        "IMPORTANT: The conditions that apply differ by position within Type-1. Not all 10 conditions apply to "
        "every position. However, 100% of applicable conditions must be met for an employee to receive any incentive. "
        "Even if you fail just ONE condition, your incentive for that month is ZERO."
    )

    add_example_box(doc, "Example: Why 100% Pass Rate Matters", [
        "Nguyen is an Assembly Inspector. She passes 9 out of 10 conditions perfectly.",
        "But she had 1 AQL failure (C5 = FAIL).",
        "Result: Nguyen receives 0 VND incentive this month.",
        "All 10 conditions must be PASS for her to receive any payment.",
    ])

    doc.add_heading("Type-2 Evaluation (Indirect Assembly QIP)", level=2)
    doc.add_paragraph(
        "For Type-2 QIP team members, performance is evaluated based on attendance conditions only (C1-C4):"
    )
    bullet(doc, "Attendance Rate >= 88%")
    bullet(doc, "Unapproved Absence <= 2 days")
    bullet(doc, "Actual Working Days > 0")
    bullet(doc, "Minimum Working Days >= 12 days")
    doc.add_paragraph(
        "The monthly incentive amount is determined by the average incentive of corresponding Type-1 positions and "
        "position-based multipliers. For leadership positions (Group Leader and above), the amount is calculated as "
        "the TYPE-1 LINE LEADER receiving average multiplied by a position-specific multiplier."
    )

    doc.add_heading("Type-3 (New QIP Members)", level=2)
    doc.add_paragraph(
        "For Type-3 QIP team members (newly hired workers with less than 30 working days), no evaluation is conducted, "
        "and the incentive policy does not apply. After completing the probation period (approximately 3 months), "
        "Type-3 members are upgraded to Type-2."
    )

    doc.add_heading("Program 2 — 6-Month Performance-Based Incentive", level=2)
    doc.add_paragraph(
        "Program 2 is a supplementary incentive program for selected QIP Talent Pool members. "
        "The purpose is to build a foundation for the yearly growing HWK QIP Talent Pool as a future HWK Quality investment, "
        "and to empower QIP manpower to achieve critical HWK-quality KPIs."
    )
    doc.add_paragraph(
        "In the V10 system, Program 2 is implemented as a Talent Pool bonus that is automatically applied alongside "
        "regular (Program 1) incentive. See Section 6 for details."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 3. SCOPE AND OBJECTS
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 3. SCOPE AND OBJECTS", level=1)

    doc.add_paragraph(
        "This policy applies to all QIP team members. "
        "Except for QIP Type-3 employees, who are newly hired workers in the process of familiarizing themselves "
        "with and training in inspection skills, with less than 30 working days."
    )

    doc.add_heading("Classification of QIP Employees", level=2)

    bold_para(doc, "Type 1 — Direct-assembly QIP: ", "anyone who inspects the final product quality")
    bullet(doc, "Inspector in Assembly line (Assembly Inspector)")
    bullet(doc, "Inspector in Repacking line")
    bullet(doc, "Inspector in AQL (AQL Inspector)")
    bullet(doc, "Trainer who inspects the inspector (Audit & Training Team)")
    bullet(doc, "QIP Model Master team")
    bullet(doc, "RQC Assembly Inspector (position code A1B — C7 and C10 exemptions apply from February 2026)")
    bullet(doc, "Line Leader")

    bold_para(doc, "Type 2 — Indirect assembly QIP: ", "anyone who contributes to the final product quality")
    bullet(doc, "Material (MTL) QIP team")
    bullet(doc, "Cutting QIP team")
    bullet(doc, "OSC/Inhouse QIP team")
    bullet(doc, "Stitching QIP team")
    bullet(doc, "Bottom QIP team")
    bullet(doc, "Packing QIP team")
    bullet(doc, "QA team")
    bullet(doc, "QIP office team")
    bullet(doc, "OCPT team")
    bullet(doc, "All other members of the QIP team who are not in Type 1 and the Testing Department (LAB)")

    bold_para(doc, "Type 3 — New QIP: ", "Newly hired workers with less than 30 working days")
    bullet(doc, "No incentive evaluation or payment during probation period")
    bullet(doc, "Automatically upgraded to Type-2 after approximately 3 months")

    doc.add_heading("Evaluation Period & Payment Schedule", level=2)
    bullet(doc, "Evaluation period: From the first day of the month to the last day of the month.")
    bullet(doc, "Reward consideration period: From the 1st to the 15th of the following month.")
    bullet(doc, "For employees still working: Payment will be made in the next payroll cycle.")
    bullet(doc, "For employees who have left: Payment will be made within 14 working days from the departure date.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 4. DEFINITIONS & ABBREVIATIONS
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 4. DEFINITIONS & ABBREVIATIONS", level=1)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Term", "Definition"])

    terms = [
        ["QIP", "Quality Improvement Program — HWK's quality-focused employee incentive system"],
        ["5Q", "Five Quality targets — HWK's strategic quality performance goals"],
        ["AQL", "Acceptable Quality Level — international sampling inspection standard. A random sample of "
                "shoes is checked; if defects exceed a limit, the entire batch fails."],
        ["5PRS", "5 Pairs Random Sampling — internal quality inspection method. 5 random pairs are checked "
                 "from each production order to verify quality."],
        ["CFA", "Certified Final Auditor — adidas-approved AQL inspection certificate"],
        ["PO", "Production Order — individual manufacturing order unit (a batch of shoes)"],
        ["C1-C10", "The 10 incentive conditions evaluated by the system (see Appendix 1 for details)"],
        ["TYPE-1", "Direct quality workers with full condition evaluation (up to 10 conditions)"],
        ["TYPE-2", "Indirect/management staff with attendance-only evaluation (4 conditions)"],
        ["TYPE-3", "New employees (< 3 months) excluded from incentive"],
        ["AR1", "Absence Record 1 — Unapproved/unauthorized absence (COUNTS against you)"],
        ["AR2", "Absence Record 2 — Approved sick leave (does NOT count against you)"],
        ["Progressive Table", "Incentive amount schedule increasing with consecutive qualifying months (150K-1M VND)"],
        ["Consecutive Months", "Number of uninterrupted months where all applicable conditions were met"],
        ["Allowance", "Exception override — manager-approved condition bypass for justified cases"],
        ["Firestore", "Google Cloud NoSQL database used for V10 data storage"],
        ["RBAC", "Role-Based Access Control — admin vs regular user permission system"],
    ]
    for t in terms:
        add_row(tbl, t)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 5. PROGRAM 1
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 5. INCENTIVE AMOUNT CALCULATION METHOD — PROGRAM 1", level=1)

    # 5.1
    doc.add_heading("5.1 The 10 Incentive Conditions", level=2)
    doc.add_paragraph(
        "The V10 system evaluates up to 10 conditions for each QIP member. Not all conditions apply to every "
        "position — the applicable conditions are determined by the Position-Condition Matrix (see Section 5.2 "
        "and Appendix 2). However, ALL applicable conditions must be met (100% pass rate) for an employee to "
        "receive incentive."
    )

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["#", "Condition", "Category", "Default Threshold", "Description"])

    conditions = [
        ["C1", "Attendance Rate", "Attendance", ">= 88%", "Attendance rate after adjusting for approved leave"],
        ["C2", "Unapproved Absence", "Attendance", "<= 2 days", "AR1 unapproved absence days only"],
        ["C3", "Actual Working Days", "Attendance", "> 0 days", "At least 1 actual working day in the month"],
        ["C4", "Minimum Working Days", "Attendance", ">= 12 days", "Must work at least 12 days in the month"],
        ["C5", "Personal AQL Failure", "AQL", "= 0 failures", "Zero personal AQL failures in current month"],
        ["C6", "Personal AQL Consecutive", "AQL", "No consecutive", "2025: 3+ months, 2026+: 2+ months"],
        ["C7", "Team/Area AQL Consecutive", "AQL", "No consecutive", "Same year-dependent rule as C6"],
        ["C8", "Area Reject Rate", "AQL", "< 3.0%", "Reject rate for assigned area below threshold"],
        ["C9", "5PRS Pass Rate", "5PRS", ">= 95%", "5PRS quality inspection pass rate"],
        ["C10", "5PRS Inspection Qty", "5PRS", ">= 100 pairs", "Minimum 5PRS inspection quantity"],
    ]
    for c in conditions:
        add_row(tbl, c)

    doc.add_paragraph()
    bold_para(doc, "Note on Configurable Thresholds: ",
              "All thresholds above are system defaults. Administrators can adjust thresholds monthly via the Admin Panel. "
              "Seven thresholds are configurable: attendance_rate, unapproved_absence, minimum_working_days, area_reject_rate, "
              "5prs_pass_rate, 5prs_min_qty, and consecutive_aql_months. All changes are logged with immutable audit trail.")

    bold_para(doc, "Note on Consecutive Failure Policy: ",
              "From 2026 onwards, the consecutive failure threshold has been updated from 3 months to 2 months. "
              "This means that 2 or more months of consecutive AQL failures will result in incentive exclusion (C6, C7). "
              "For 2025, only 3 or more months of consecutive failures cause exclusion.")

    # 5.2
    doc.add_heading("5.2 Position-Condition Application Matrix (Type-1)", level=2)
    doc.add_paragraph(
        "Different positions within Type-1 have different applicable conditions. "
        "The following table summarizes which conditions apply to each position:"
    )

    tbl = doc.add_table(rows=1, cols=11)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Position", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9", "C10"])

    matrix = [
        ["ASSEMBLY INSPECTOR", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "Yes"],
        ["RQC ASSEMBLY INSPECTOR\n(A1B, from 2026.02)", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "-"],
        ["AQL INSPECTOR", "Yes", "Yes", "Yes", "Yes", "Yes", "-", "-", "-", "-", "-"],
        ["LINE LEADER", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "-", "-", "-"],
        ["AUDITOR & TRAINER", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "Yes", "-", "-"],
        ["MODEL MASTER", "Yes", "Yes", "Yes", "Yes", "-", "-", "-", "Yes", "-", "-"],
        ["Management (Group Leader,\nSupervisor, Manager, etc.)", "Yes", "Yes", "Yes", "Yes", "-", "-", "-", "-", "-", "-"],
    ]
    for m in matrix:
        add_row(tbl, m)

    doc.add_paragraph()
    bold_para(doc, "RQC Assembly Inspector (A1B): ",
              "From February 2026, position code A1B is classified as RQC_ASSEMBLY_INSPECTOR. "
              "C7 (Team/Area AQL) and C10 (5PRS Inspection Qty) are exempted because RQC inspectors perform "
              "process checking and reports, not line inspection. Before February 2026, A1B is treated as regular Assembly Inspector.")

    # 5.3 — Calculation Examples
    doc.add_heading("5.3 Attendance Calculation — Step-by-Step Examples", level=2)
    doc.add_paragraph(
        "Understanding how attendance is calculated is critical. Here is the exact formula the system uses:"
    )

    bold_para(doc, "Attendance Rate Formula:")
    bullet(doc, "Step 1: Calculate Absence Days = Total Working Days - Actual Working Days - Approved Leave Days")
    bullet(doc, "Step 2: Calculate Absence Rate = (Absence Days / Total Working Days) x 100")
    bullet(doc, "Step 3: Calculate Attendance Rate = 100 - Absence Rate")
    doc.add_paragraph(
        "Key point: Approved leave (annual leave, maternity leave, sick leave AR2, business trip, etc.) "
        "does NOT count as absence. Only AR1 (unauthorized absence) counts as absence."
    )

    add_example_box(doc, "Example 1: Employee PASSES Attendance (C1)", [
        "Tran works in February 2026. The month has 27 total working days.",
        "- She actually worked 23 days",
        "- She took 2 days of approved annual leave (Phep nam)",
        "",
        "Step 1: Absence Days = 27 - 23 - 2 = 2 days",
        "Step 2: Absence Rate = (2 / 27) x 100 = 7.4%",
        "Step 3: Attendance Rate = 100 - 7.4 = 92.6%",
        "",
        "Result: 92.6% >= 88% threshold --> C1 = PASS",
    ])

    add_example_box(doc, "Example 2: Employee FAILS Attendance (C1)", [
        "Le works in a month with 26 total working days.",
        "- He actually worked 18 days",
        "- He took 1 day of approved sick leave (AR2)",
        "- He was absent 7 days (some approved, some not)",
        "",
        "Step 1: Absence Days = 26 - 18 - 1 = 7 days",
        "Step 2: Absence Rate = (7 / 26) x 100 = 26.9%",
        "Step 3: Attendance Rate = 100 - 26.9 = 73.1%",
        "",
        "Result: 73.1% < 88% threshold --> C1 = FAIL --> No incentive this month",
    ])

    add_example_box(doc, "Example 3: Maternity Leave Does NOT Hurt You", [
        "Hoa takes 5 days of maternity leave (Sinh thuong 1 con) in a 27-day month.",
        "She works the other 22 days.",
        "",
        "Step 1: Absence Days = 27 - 22 - 5 = 0 days",
        "Step 2: Absence Rate = (0 / 27) x 100 = 0%",
        "Step 3: Attendance Rate = 100 - 0 = 100%",
        "",
        "Result: 100% >= 88% --> C1 = PASS",
        "Maternity leave is approved leave, so it does not count as absence.",
    ])

    doc.add_paragraph()

    # Absence Classification Table
    doc.add_heading("5.3.1 Absence Classification — What Counts and What Doesn't", level=3)
    doc.add_paragraph(
        "This is one of the most important things to understand. Not all types of leave are treated the same:"
    )

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Leave Type (Vietnamese)", "English Translation", "Counts as Absence?"], bg="2F5496")

    # Approved leaves (NOT counted)
    approved_leaves = [
        ["Sinh thuong 1 con", "Maternity Leave", "NO (Approved)"],
        ["Phep nam", "Annual Leave", "NO (Approved)"],
        ["Vang co phep", "Approved Absence", "NO (Approved)"],
        ["Duong suc sinh thuong", "Postpartum Rest", "NO (Approved)"],
        ["Kham thai binh thuong", "Prenatal Checkup", "NO (Approved)"],
        ["Con duoi 3 tuoi bi benh", "Childcare Leave (child < 3 years)", "NO (Approved)"],
        ["AR2 - om ngan ngay", "Short Sick Leave (AR2)", "NO (Approved)"],
        ["Di cong tac", "Business Trip", "NO (Approved)"],
        ["Nghia vu quan su", "Military Service", "NO (Approved)"],
        ["Di lam khong quet the", "Card Not Swiped (worked but forgot card)", "NO (Approved)"],
        ["Cong nhan vien moi", "New Employee Exception", "NO (Approved)"],
        ["Nghi bu", "Compensatory Leave", "NO (Approved)"],
    ]
    for lv in approved_leaves:
        add_row(tbl, lv, bg_color="E8F5E9")  # green tint

    # Unapproved (COUNTED)
    unapproved_leaves = [
        ["AR1 - Vang khong phep", "Unauthorized Absence", "YES (Unapproved - AR1)"],
        ["AR1 - Gui thu", "Written Notice Absence", "YES (Unapproved - AR1)"],
    ]
    for lv in unapproved_leaves:
        add_row(tbl, lv, bg_color="FFEBEE")  # red tint

    doc.add_paragraph()
    bold_para(doc, "Simple rule: ",
              "If the leave type starts with \"AR1\", it COUNTS as absence and hurts your attendance rate. "
              "Everything else is approved leave and does NOT count against you.")

    doc.add_page_break()

    # 5.4 — Progressive Table
    doc.add_heading("5.4 Type-1 Progressive Incentive Table", level=2)
    doc.add_paragraph(
        "For Type-1 members who pass all applicable conditions, the incentive amount increases with continuous "
        "months of qualifying. The longer your streak of passing all conditions, the more you earn:"
    )

    # Horizontal table
    tbl = doc.add_table(rows=2, cols=14)
    tbl.style = 'Table Grid'
    headers = ["Months", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12-15"]
    amounts = ["VND (x1,000)", "150", "250", "300", "350", "400", "450", "500", "650", "750", "850", "950", "1,000"]

    for i, text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")

    for i, text in enumerate(amounts):
        cell = tbl.rows[1].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8)
        if i == 0:
            run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Key rules for the progressive table:")
    bullet(doc, "Maximum cap: 1,000,000 VND per month (reached at 12 continuous months)")
    bullet(doc, "Reset on failure: If an employee fails to meet conditions in any month, continuous months reset to 0")
    bullet(doc, "Carry over: Continuous months carry over from the previous month upon meeting conditions")
    bullet(doc, "Maximum tracking: 15 months (amounts remain at 1,000,000 VND from month 12 onwards)")

    add_example_box(doc, "Example: How the Progressive Table Works", [
        "Month 1 (Jan): All conditions PASS -> Consecutive month = 1 -> 150,000 VND",
        "Month 2 (Feb): All conditions PASS -> Consecutive month = 2 -> 250,000 VND",
        "Month 3 (Mar): All conditions PASS -> Consecutive month = 3 -> 300,000 VND",
        "Month 4 (Apr): C5 FAIL (AQL failure) -> Consecutive months RESET to 0 -> 0 VND",
        "Month 5 (May): All conditions PASS -> Consecutive month = 1 (restart!) -> 150,000 VND",
        "",
        "As you can see, one failure resets your streak back to month 1.",
        "This is why consistent quality performance is so important!",
    ])

    # 5.5 — AQL Inspector 3-Part
    doc.add_heading("5.5 AQL Inspector — 3-Part Incentive Calculation", level=2)
    doc.add_paragraph(
        "AQL Inspectors receive a special 3-part incentive calculation. The final monthly amount is the sum of all three parts:"
    )

    bold_para(doc, "Part 1: AQL Inspection Evaluation Result")
    doc.add_paragraph(
        "Based on the progressive table (same as Section 5.4 above). Requires rejection rate by adidas/T1QM/3rd Party "
        "inspectors < 3%."
    )

    bold_para(doc, "Part 2: CFA Certificate Incentive")
    doc.add_paragraph(
        "AQL Inspectors with a valid adidas-approved AQL certificate (CFA) receive a fixed monthly amount of 700,000 VND. "
        "This is paid every month as long as the certificate remains valid and the inspector meets all applicable conditions."
    )

    bold_para(doc, "Part 3: HWK Claim Prevention Incentive")
    doc.add_paragraph(
        "Considering the role of AQL inspectors as the final quality gatekeeper, an additional incentive for "
        "preventing HWK complaints is provided. This amount increases progressively starting from month 4:"
    )

    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = 'Table Grid'
    headers2 = ["Months", "1-3", "4-6", "7-9", "10-12", "13-15"]
    amounts2 = ["VND (x1,000)", "0", "300", "500", "700", "900"]
    for i, text in enumerate(headers2):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")
    for i, text in enumerate(amounts2):
        cell = tbl.rows[1].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True

    doc.add_paragraph()

    add_example_box(doc, "Example: AQL Inspector Total Incentive at Different Stages", [
        "--- At Month 3 (with CFA certificate): ---",
        "  Part 1 (Progressive): 300,000 VND",
        "  Part 2 (CFA Certificate): 700,000 VND",
        "  Part 3 (HWK Prevention): 0 VND (months 1-3 = 0)",
        "  Total: 1,000,000 VND",
        "",
        "--- At Month 7 (with CFA certificate): ---",
        "  Part 1 (Progressive): 500,000 VND",
        "  Part 2 (CFA Certificate): 700,000 VND",
        "  Part 3 (HWK Prevention): 500,000 VND",
        "  Total: 1,700,000 VND",
        "",
        "--- At Month 12 (with CFA certificate): ---",
        "  Part 1 (Progressive): 1,000,000 VND",
        "  Part 2 (CFA Certificate): 700,000 VND",
        "  Part 3 (HWK Prevention): 700,000 VND",
        "  Total: 2,400,000 VND",
    ])

    doc.add_page_break()

    # 5.6
    doc.add_heading("5.6 Line Leader Incentive Calculation (Type-1)", level=2)
    doc.add_paragraph(
        "Line Leaders in Type-1 receive incentive calculated based on their subordinate team's performance:"
    )
    bold_para(doc, "Formula: ", "Subordinate Total Incentive x 12%")
    doc.add_paragraph(
        "The system uses a subordinate mapping to identify all team members under each Line Leader. "
        "Only subordinates who received incentive (> 0 VND) are counted in the total. "
        "Line Leaders must also pass all applicable conditions (C1, C2, C3, C4, C7) to receive the incentive."
    )

    add_example_box(doc, "Example: Line Leader Incentive", [
        "Line Leader Minh has 8 inspectors in his team.",
        "- 6 inspectors passed all conditions and received incentive:",
        "  300,000 + 250,000 + 400,000 + 150,000 + 500,000 + 350,000 = 1,950,000 VND total",
        "- 2 inspectors failed and received 0 VND",
        "",
        "Minh's incentive = 1,950,000 x 12% = 234,000 VND",
        "(Only if Minh himself passes C1, C2, C3, C4, and C7)",
    ])

    doc.add_paragraph("Line Leader incentive exclusion cases:")
    bullet(doc, "If any inspector in the Line Leader's team has consecutive AQL failures exceeding the threshold "
           "(C7), the Line Leader will not receive monthly incentive.")

    # 5.7
    doc.add_heading("5.7 Type-2 Incentive Calculation", level=2)
    doc.add_paragraph(
        "Type-2 members must meet attendance conditions (C1-C4) at 100% pass rate. The incentive amount is "
        "determined based on the average incentive of corresponding Type-1 positions."
    )

    bold_para(doc, "TYPE-2 Position Mapping for Incentive Calculation:")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Type-2 Position", "Incentive Reference", "Description"])

    mappings = [
        ["BOTTOM INSPECTOR", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["CUTTING INSPECTOR", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["MTL INSPECTOR", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["STITCHING INSPECTOR", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["OSC INSPECTOR", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["OCPT STAFF", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["RQC", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["AQL INSPECTOR (TYPE-2)", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["QA TEAM (default)", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["QA TEAM (code QA3A)", "GROUP LEADER average", "Special: mapped to Group Leader average"],
        ["QA TEAM (code QA3B)", "ASSEMBLY INSPECTOR average", "Uses Type-1 Assembly Inspector receiving average"],
        ["LINE LEADER (TYPE-2)", "LINE LEADER average", "Uses Type-1 Line Leader receiving average"],
    ]
    for m in mappings:
        add_row(tbl, m)

    add_example_box(doc, "Example: Type-2 Inspector Incentive", [
        "Loan is a Cutting Inspector (Type-2). She passes all 4 attendance conditions.",
        "This month, the average incentive of Type-1 Assembly Inspectors who received incentive is 320,000 VND.",
        "",
        "Loan's incentive = 320,000 VND (same as the Assembly Inspector average)",
    ])

    # 5.8
    doc.add_heading("5.8 Type-2 Leadership Position Multipliers", level=2)
    doc.add_paragraph(
        "For mid-to-leadership management positions in Type-2, the incentive is calculated as the TYPE-1 LINE LEADER "
        "receiving average multiplied by a position-specific multiplier:"
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Position", "Multiplier", "Formula", "Description"])

    multipliers = [
        ["GROUP LEADER", "2.0x", "LL Avg x 2.0", "Line Leader receiving average x 2.0"],
        ["SUPERVISOR /\n(V) SUPERVISOR /\nA.SUPERVISOR", "2.5x", "LL Avg x 2.5", "Line Leader receiving average x 2.5"],
        ["A.MANAGER", "3.0x", "LL Avg x 3.0", "Line Leader receiving average x 3.0"],
        ["MANAGER", "3.5x", "LL Avg x 3.5", "Line Leader receiving average x 3.5"],
        ["S.MANAGER", "4.0x", "LL Avg x 4.0", "Line Leader receiving average x 4.0"],
    ]
    for m in multipliers:
        add_row(tbl, m)

    doc.add_paragraph()

    add_example_box(doc, "Example: Leadership Multiplier Calculation", [
        "This month, TYPE-1 Line Leaders who received incentive have an average of 400,000 VND.",
        "",
        "Group Leader (Type-2) = 400,000 x 2.0 = 800,000 VND",
        "(V) Supervisor (Type-2) = 400,000 x 2.5 = 1,000,000 VND",
        "A.Manager (Type-2) = 400,000 x 3.0 = 1,200,000 VND",
        "Manager (Type-2) = 400,000 x 3.5 = 1,400,000 VND",
        "S.Manager (Type-2) = 400,000 x 4.0 = 1,600,000 VND",
    ])

    bold_para(doc, "Note: ",
              "\"LINE LEADER receiving average\" means the average of Type-1 Line Leader incentives for recipients only "
              "(employees with incentive > 0 VND). Employees with 0 VND are excluded from the average calculation.")

    # 5.9
    doc.add_heading("5.9 Audit & Training Team (Trainer) Incentive", level=2)
    doc.add_paragraph(
        "Each training staff member has their own area of responsibility for employee inspection and training."
    )
    bullet(doc, "If the error rate (Area Reject Rate, C8) in the area they are responsible for is below threshold: Meets the requirements.")
    bullet(doc, "If the error rate is above threshold: Does not meet the requirements.")
    bullet(doc, "Applicable conditions: C1, C2, C3, C4, C7 (Team/Area AQL Consecutive), C8 (Area Reject Rate)")
    doc.add_paragraph(
        "Each Auditor/Trainer is assigned to specific building areas (A, B, C, D, or Repacking). "
        "The area mapping is maintained in the system configuration. "
        "The progressive table (Section 5.4) applies for incentive amount calculation."
    )

    # 5.10
    doc.add_heading("5.10 Model Master Team Incentive", level=2)
    doc.add_paragraph(
        "The Model Master team's incentive is based on the HWK AQL rejection percentage across all factories."
    )
    bullet(doc, "If the official HWK AQL rejection percentage across all factories (Area Reject Rate, C8) is below threshold: eligible.")
    bullet(doc, "Applicable conditions: C1, C2, C3, C4, C8 (Area Reject Rate)")
    doc.add_paragraph("The progressive table (Section 5.4) applies for incentive amount calculation.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 6. PROGRAM 2
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 6. INCENTIVE AMOUNT CALCULATION METHOD — PROGRAM 2", level=1)

    doc.add_paragraph(
        'Program 2 is the "6-Month-Performance-Based-Incentive Program." Only selected QIP team members are '
        "evaluated through a comprehensive process spanning six months. The evaluation method or focus may vary "
        "every six months based on the quality team's targets and HWK quality status."
    )

    bold_para(doc, "IMPORTANT: ", "Anyone who receives incentive based on Program 1 can also receive incentive based on Program 2.")

    doc.add_paragraph()
    doc.add_paragraph(
        "The following 6-step process is used for Program 2 management. "
        "In the V10 system, the Talent Pool bonus is implemented as an automatic supplementary incentive "
        "managed through the system configuration."
    )

    steps = [
        ("Step 1: Target Set-Up", "Every January and July, the QIP team reviews the previous six months of quality performance "
         "and plans for the next six months. QIP managers and leadership team select critical quality topics, define measurable KPIs, "
         "or initiate a quality project operation plan."),
        ("Step 2: Task Assignment", "The QIP managers and leadership team select the \"right person\" for each target, regardless of position. "
         "The incentive amount and payment period are defined. Approval is obtained from the COO."),
        ("Step 3: Monthly Monitoring & Feedback", "The QIP QA team, office team, and managers conduct monthly reviews using measurable KPIs "
         "and provide feedback to all candidate members of the HWK QIP Talent Pool."),
        ("Step 4: Final Evaluation", "After six months, QIP managers review the plan and performance outcomes, making a final decision on "
         "who will receive certification for the HWK QIP Talent Pool."),
        ("Step 5: Confirmation of Talent Pool Membership", "The QIP manager requests approval from HWK HR team and HWK COO via the HS "
         "groupware system. Once approved, the updated talent pool information is integrated into the system configuration."),
        ("Step 6: Developing the HWK QIP Talent Pool", "The process is continuous with a new cycle every 6 months. "
         "The goal is to certify at least 3-5% of total QIP team members as talented individuals over time. "
         "Certification remains valid for 12 months."),
    ]
    for title, desc in steps:
        bold_para(doc, title)
        doc.add_paragraph(desc)

    doc.add_paragraph()
    bold_para(doc, "V10 System Implementation:")
    bullet(doc, "Auto-apply: Talent Pool bonus is automatically applied alongside regular incentive")
    bullet(doc, "Stackable: Talent Pool bonus stacks with regular (Program 1) incentive")
    bullet(doc, "Payment timing: Together with regular monthly incentive cycle")
    bullet(doc, "Conditions requirement: No — Talent Pool members receive bonus regardless of condition pass/fail")
    bullet(doc, "Exclusion: Resigned employees are automatically excluded")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 7. EXCEPTIONAL CLAUSES
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 7. SUMMARY OF EXCEPTIONAL CLAUSES", level=1)

    doc.add_heading("7.1 Conditions for NOT Receiving Incentive", level=2)

    bold_para(doc, "For Type-1 (Program 1):")
    bullet(doc, "If any applicable condition among the 10 conditions is not met (100% pass required), monthly incentive is not provided.")
    bullet(doc, "If attendance rate < threshold, if unapproved absence > threshold, if actual working days = 0, or if minimum working days < threshold.")
    bullet(doc, "If 5PRS pass rate < threshold (for positions where C9 applies), monthly incentive is not provided.")
    bullet(doc, "If 5PRS inspection quantity < threshold (for positions where C10 applies), monthly incentive is not provided.")
    bullet(doc, "If there is an adidas official claim and the responsibility belongs to the individual, monthly incentive is not provided.")
    bullet(doc, "If the number of personal AQL PO reject is more than 0 (for positions where C5 applies), monthly incentive is not provided.")
    bullet(doc, "For Line Leaders: If any inspector in the team has consecutive AQL failures exceeding the threshold, the Line Leader will not receive monthly incentive.")
    bullet(doc, "For Trainers: If the assigned building's official monthly AQL reject rate exceeds the area reject threshold, monthly incentive is not provided.")
    bullet(doc, "For Trainers: If any inspector in the trainer's responsible building has consecutive AQL failures exceeding the threshold, the trainer will not receive monthly incentive.")

    bold_para(doc, "For Type-2 (Program 1):")
    bullet(doc, "If attendance conditions (C1-C4) are not all met at 100%, monthly incentive is not provided.")
    bullet(doc, "If there is an adidas official claim or critical internal complaint and the responsibility belongs to the individual, monthly incentive is not provided.")

    bold_para(doc, "For Type-3 (Program 1):")
    bullet(doc, "For ALL QIP members of Type-3, monthly incentive is not provided (policy exclusion).")

    bold_para(doc, "For Program 2:")
    bullet(doc, "If the QIP manager's suggested targets, person in charge, incentive amount, and provision period are rejected by COO during the 6-month renewal, all Program 2 processes are not effective.")

    doc.add_heading("7.2 Exceptional Cases for RECEIVING Incentive (Manager Override / Allowance)", level=2)

    doc.add_paragraph(
        "The V10 system includes an Allowance/Exception Override system accessible through the Admin Dashboard. "
        "This allows authorized managers to grant exceptions in the following cases:"
    )

    bullet(doc, "Due to limited annual leave days or unexpected urgent personal issues, if anyone cannot meet the attendance "
           "prerequisite, QIP managers can review and validate the case as \"exceptional.\" Without counting the exceptional "
           "case into absent days, if the employee can meet the attendance rule, monthly incentive is provided.")
    bullet(doc, "If an AQL reject PO case is related to \"not inspector's poor performance\" but \"production team's poor quality "
           "ownership issue\" (such as intended wrong packing finding from AQL room), the AQL reject PO will not count for the "
           "related Inspector's monthly AQL performance validation.")
    bullet(doc, "If 5PRS validation result is under threshold and total validation quantity does not meet expectation, QIP manager "
           "can review case by case and make an exceptional allowance if the root cause is reasonable (such as working area rotation, R&R update, etc.).")
    bullet(doc, "Case by case, exceptional cases can be allowed by QIP manager to provide incentive to a right-performance-made "
           "member, consistent with the concept of providing incentive for QIP team motivation and fostering quality culture.")

    bold_para(doc, "V10 System Allowance Implementation:")
    doc.add_paragraph("The system provides 4 reason codes for documenting allowance overrides:")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Reason Code", "Description"])
    codes = [
        ["MEDICAL", "Medical reason (illness, injury, hospitalization)"],
        ["COMPANY_ORDER", "Company order (reassignment, mandatory tasks)"],
        ["NATURAL_DISASTER", "Natural disaster (flooding, typhoon, etc.)"],
        ["OTHER", "Other documented reason (free text justification required)"],
    ]
    for c in codes:
        add_row(tbl, c)

    doc.add_paragraph()
    doc.add_paragraph("Allowance workflow in the system:")
    bullet(doc, "1. Administrator searches for the employee in the Allowance tab")
    bullet(doc, "2. System displays current condition status (pass/fail for each applicable condition)")
    bullet(doc, "3. Administrator selects failed conditions to override")
    bullet(doc, "4. Administrator provides reason code + detailed justification (required)")
    bullet(doc, "5. Preview screen shows the override impact before applying")
    bullet(doc, "6. On Apply: system updates employee data, recalculates incentive amount, updates dashboard summary")
    bullet(doc, "7. Allowance can be revoked later (restores original condition status)")
    bullet(doc, "8. Bulk recalculation available for all allowance employees with 0 incentive")

    doc.add_paragraph()
    doc.add_paragraph(
        "All allowance overrides are tracked with audit trails including: override date, authorized manager (email), "
        "affected employee, original condition status, overridden conditions, override reason code, reason detail, "
        "and recalculated incentive amount."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 8. SYSTEM ARCHITECTURE & DASHBOARD
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 8. SYSTEM ARCHITECTURE & DASHBOARD (V10)", level=1)

    doc.add_paragraph(
        "The HWK QIP Incentive System V10 is a cloud-based platform that automates the entire incentive calculation "
        "pipeline, from data collection to final payment determination. All employee data is loaded securely from "
        "Firestore after authentication, replacing the previous approach of embedding data directly in HTML files."
    )
    doc.add_paragraph(
        "You can access the dashboard at: https://moonkaicuzui.github.io/hwk-qip-incentive-v10/"
    )

    doc.add_heading("8.1 How to Use the Dashboard — Step by Step", level=2)

    doc.add_paragraph(
        "Here is a step-by-step guide to accessing and using the QIP Incentive Dashboard:"
    )

    bold_para(doc, "Step 1: Log In")
    doc.add_paragraph(
        "Open the dashboard URL in your browser. You will see a login screen. "
        "Enter your email address and password (provided by your administrator). "
        "Click the Login button."
    )

    bold_para(doc, "Step 2: Select Month and Year")
    doc.add_paragraph(
        "After logging in, you will see the Month/Year selector screen. "
        "Choose the month and year you want to view, then click the \"Go to Dashboard\" button."
    )
    add_screenshot(doc, "02_selector_en.png",
                   "Figure 1: Month/Year Selection Screen")

    bold_para(doc, "Step 3: View the Dashboard Summary")
    doc.add_paragraph(
        "The Summary tab shows the overall KPI cards at the top: total recipients, payment rate, "
        "total incentive amount, and trend indicators compared to the previous month. Below the KPI cards, "
        "you can see TYPE distribution tables, condition pass/fail charts, and trend charts."
    )
    add_screenshot(doc, "03_dashboard_summary_en.png",
                   "Figure 2: Dashboard Summary Tab — KPI cards and overview")

    bold_para(doc, "Step 4: Check Your Individual Status")
    doc.add_paragraph(
        "Click on the \"Individual Detail\" tab. You can search for any employee by name or ID. "
        "The list shows each person's condition pass/fail status and incentive amount."
    )
    add_screenshot(doc, "04_individual_en.png",
                   "Figure 3: Individual Detail Tab — searchable employee list")

    bold_para(doc, "Step 5: View Employee Detail Modal")
    doc.add_paragraph(
        "Click on any employee's row to open the detail modal. This shows ALL applicable conditions "
        "for that employee, whether each condition passed or failed, the actual values, and the final "
        "incentive calculation breakdown."
    )
    add_screenshot(doc, "05_employee_modal_en.png",
                   "Figure 4: Employee Detail Modal — condition breakdown")

    bold_para(doc, "Step 6: Review Incentive Criteria")
    doc.add_paragraph(
        "The \"Incentive Criteria\" tab shows the current threshold values for all 10 conditions, "
        "along with descriptions of what each condition means. This is useful as a quick reference "
        "to understand what targets you need to meet."
    )
    add_screenshot(doc, "06_criteria_en.png",
                   "Figure 5: Incentive Criteria Tab — current thresholds and descriptions")

    doc.add_page_break()

    doc.add_heading("8.2 Data Flow & CI/CD Pipeline", level=2)
    doc.add_paragraph("The system follows this automated data flow (runs every 6 hours + manual trigger):")
    bullet(doc, "Google Drive: Source data files (basic manpower, attendance, AQL reports, 5PRS data) are uploaded to Google Drive.")
    bullet(doc, "GitHub Actions: Automated CI/CD pipeline downloads data from Google Drive, syncs thresholds and configs from Firestore.")
    bullet(doc, "Auto-Detection: System automatically detects stale months requiring recalculation.")
    bullet(doc, "Attendance Conversion: Raw attendance data is converted and schema-validated.")
    bullet(doc, "Python Calculation Engine: Processes all 10 conditions for each employee and calculates incentive amounts.")
    bullet(doc, "Firestore Upload: Calculated results are uploaded to Google Cloud Firestore database.")
    bullet(doc, "Email Processing: Pending email notifications are sent via SMTP (mail.hsvina.com).")
    bullet(doc, "Deployment: Dashboard is deployed to GitHub Pages.")

    doc.add_heading("8.3 Dashboard Features (8 Tabs)", level=2)
    bullet(doc, "Summary Tab: KPI cards (recipients count, payment rate, total amount with trend indicators), TYPE distribution table, "
           "condition pass/fail charts, trend chart, talent pool display.")
    bullet(doc, "Position Detail Tab: Position-specific breakdown tables with condition status.")
    bullet(doc, "Individual Detail Tab: Per-employee searchable/filterable list with detailed condition modal (250ms debounce).")
    bullet(doc, "Incentive Criteria Tab: Current threshold reference and condition descriptions.")
    bullet(doc, "Organization Chart Tab: QIP team hierarchical structure with expected incentive calculations.")
    bullet(doc, "Team Management Tab: Team-level analytics and performance comparison.")
    bullet(doc, "Validation Summary Tab: 12 KPI validation cards, data integrity checks, system verification.")
    bullet(doc, "Attendance Lookup Tab: Individual attendance calendar with daily status lookup.")

    doc.add_heading("8.4 Admin Panel (5 Tabs)", level=2)
    doc.add_paragraph("The Admin Panel is accessible to administrators only (role-based access control):")
    bullet(doc, "Thresholds: Adjust 7 configurable thresholds per month with change history and audit trail.")
    bullet(doc, "Config Management: Position mapping, progressive table configuration, type classification rules.")
    bullet(doc, "System: Pipeline status monitoring, system configuration, admin email list management.")
    bullet(doc, "Data Lookup: Direct Firestore data querying and verification.")
    bullet(doc, "Allowances: Exception override management with apply/revoke workflow and bulk recalculation.")

    doc.add_heading("8.5 Threshold Management", level=2)
    doc.add_paragraph(
        "All threshold values are centrally managed in Firestore and can be adjusted by administrators through the Admin Dashboard. "
        "Every change creates an immutable record in the threshold_history collection including: changed fields, old values, "
        "new values, changed by (email), and timestamp. The Python calculation engine syncs these values before each run."
    )

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Parameter", "System Default", "Configurable"])

    params = [
        ["Attendance Rate (C1)", "88%", "Yes"],
        ["Unapproved Absence (C2)", "2 days", "Yes"],
        ["Minimum Working Days (C4)", "12 days", "Yes"],
        ["Area Reject Rate (C8)", "3.0%", "Yes"],
        ["5PRS Pass Rate (C9)", "95%", "Yes"],
        ["5PRS Min Quantity (C10)", "100 pairs", "Yes"],
        ["Consecutive AQL Months (C6, C7)", "3 months", "Yes"],
    ]
    for p_row in params:
        add_row(tbl, p_row)

    doc.add_heading("8.6 Feedback System", level=2)
    doc.add_paragraph(
        "The system includes a built-in feedback/issue tracking system accessible to all authenticated users:"
    )
    bullet(doc, "Feedback Types: BUG, IMPROVEMENT, NEW_FEATURE, UI_UX, DATA, OTHER")
    bullet(doc, "Priorities: Critical, High, Medium, Low")
    bullet(doc, "Status Flow: SUBMITTED -> REVIEWING -> IN_PROGRESS -> COMPLETED (or REJECTED)")
    bullet(doc, "Features: Image attachments (up to 3), multiple notification recipients, admin reply with email notification")

    doc.add_heading("8.7 Email Notification System", level=2)
    doc.add_paragraph("Three automated email notifications via Cloud Functions (Node.js 22, asia-northeast3):")
    bullet(doc, "New feedback submitted -> All administrators receive email notification (3-language: ko/en/vi)")
    bullet(doc, "Feedback status changed -> Feedback reporter receives status update notification")
    bullet(doc, "Admin replies to feedback -> Feedback reporter receives reply email")
    doc.add_paragraph("SMTP server: mail.hsvina.com:465 (SSL). All emails logged to email_logs collection for audit.")

    doc.add_heading("8.8 Multi-Language Support", level=2)
    doc.add_paragraph("The entire dashboard supports three languages: Korean (ko), English (en), and Vietnamese (vi). "
                      "All UI strings, error messages, and labels are managed through the i18n system.")

    doc.add_heading("8.9 Authentication & RBAC", level=2)
    doc.add_paragraph("Firebase Email/Password authentication with Role-Based Access Control:")
    bullet(doc, "Regular users: Read-only dashboard access + feedback submission")
    bullet(doc, "Administrators: Full access including admin panel, threshold management, allowance overrides, feedback management")
    doc.add_paragraph("Admin email list is dynamically managed in Firestore (system/config.admin_emails).")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 9. QUICK REFERENCE & FAQ
    # ══════════════════════════════════════════
    doc.add_heading("SECTION 9. QUICK REFERENCE & FAQ", level=1)

    doc.add_heading("9.1 Quick Reference Card", level=2)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Question", "Answer", "Details"])

    qr_items = [
        ["Who gets incentive?", "All QIP members except Type-3", "Type-1: up to 10 conditions\nType-2: 4 attendance conditions"],
        ["How much can I earn?", "150,000 - 1,000,000 VND/month", "Increases with consecutive months\nAQL Inspectors: up to 2,400,000"],
        ["What if I fail 1 condition?", "0 VND for that month", "100% pass rate required\nConsecutive months reset to 0"],
        ["Does annual leave hurt me?", "NO", "Approved leave does NOT count as absence"],
        ["Does sick leave (AR2) hurt?", "NO", "AR2 is approved leave"],
        ["What hurts my attendance?", "Only AR1 (unauthorized absence)", "AR1 = Vang khong phep or Gui thu"],
        ["What is the maximum cap?", "1,000,000 VND/month", "Reached at 12 consecutive months"],
        ["Can exceptions be made?", "Yes — Allowance system", "Manager can override failed conditions\nwith documented justification"],
        ["How do I check my status?", "Dashboard website", "Login -> Select month -> Individual tab"],
        ["When is payment?", "Next payroll cycle", "After reward consideration (1st-15th)"],
    ]
    for item in qr_items:
        add_row(tbl, item)

    doc.add_heading("9.2 Frequently Asked Questions", level=2)

    faq_items = [
        ("Q: I took annual leave for 3 days. Will this affect my incentive?",
         "A: No. Annual leave (Phep nam) is approved leave and does NOT count as absence. "
         "Your attendance rate calculation will subtract these days from the absence calculation. "
         "Example: 27 working days, worked 22 days, 3 days annual leave = Absence is 27-22-3 = 2 days, "
         "Attendance Rate = 100 - (2/27 x 100) = 92.6% which is above the 88% threshold."),

        ("Q: I was sick for 2 days and got a doctor's note (AR2). Does this hurt me?",
         "A: No. AR2 (sick leave with documentation) is approved leave. It does not count as absence. "
         "However, if you were absent WITHOUT a doctor's note, it would be classified as AR1 "
         "(unauthorized absence), which DOES count against you."),

        ("Q: I failed one AQL inspection this month. What happens?",
         "A: If C5 (Personal AQL Failure) applies to your position, even 1 AQL failure means C5 = FAIL. "
         "Since 100% pass rate is required, your incentive for this month is 0 VND, and your consecutive "
         "months counter resets to 0. Next month, if you pass everything, you start from month 1 again (150,000 VND)."),

        ("Q: I had 11 consecutive months of passing. I failed this month. What happens next month?",
         "A: Your consecutive months reset to 0. If you pass all conditions next month, you start from "
         "month 1 (150,000 VND). You would need to build up 12 consecutive months again to reach the "
         "maximum of 1,000,000 VND."),

        ("Q: I am a new employee. When do I start receiving incentive?",
         "A: New employees are classified as Type-3 for approximately 3 months (until they reach 30 working days). "
         "During Type-3, no incentive is paid. After the probation period, you are upgraded to Type-2 and will "
         "receive attendance-based incentive."),

        ("Q: My manager says my AQL failure was not my fault. Can I still get incentive?",
         "A: Yes, possibly. Your manager can use the Allowance system to override the failed condition. "
         "They must provide a reason code and detailed justification. The system will then recalculate "
         "your incentive with the overridden condition marked as PASS."),

        ("Q: What is the difference between Type-1 and Type-2?",
         "A: Type-1 workers directly inspect final product quality (assembly inspectors, AQL inspectors, etc.) "
         "and are evaluated on up to 10 conditions. Type-2 workers contribute to quality indirectly (cutting, "
         "stitching, material inspection, etc.) and are evaluated only on 4 attendance conditions."),

        ("Q: Can I receive both Program 1 and Program 2 incentive?",
         "A: Yes! If you are selected for the Talent Pool (Program 2), you receive the Talent Pool bonus "
         "IN ADDITION to your regular Program 1 incentive. They stack together."),

        ("Q: What does 'consecutive months' mean exactly?",
         "A: It means the number of months IN A ROW where you passed all applicable conditions and received incentive. "
         "If you pass in January, February, and March, your consecutive months are 3. If you fail in April, "
         "it resets to 0. If you pass again in May, it starts at 1."),

        ("Q: I am an RQC Assembly Inspector (A1B). Which conditions apply to me?",
         "A: From February 2026, RQC Assembly Inspectors have these conditions: C1 (Attendance Rate), "
         "C2 (Unapproved Absence), C3 (Actual Working Days), C4 (Minimum Working Days), C5 (Personal AQL Failure), "
         "C6 (Personal AQL Consecutive), and C9 (5PRS Pass Rate). You are EXEMPT from C7 and C10."),
    ]

    for question, answer in faq_items:
        bold_para(doc, question)
        doc.add_paragraph(answer)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════
    # APPENDIX 1 — DETAILED CONDITION GUIDE
    # ══════════════════════════════════════════
    doc.add_heading("APPENDIX 1: 10 INCENTIVE CONDITIONS — DETAILED GUIDE WITH EXAMPLES", level=1)

    doc.add_paragraph(
        "This appendix explains each of the 10 incentive conditions in detail. For each condition, you will find: "
        "what it measures, why it matters, how it is calculated, which positions it applies to, and real-world examples."
    )

    # ── C1 ──
    doc.add_heading("C1 — Attendance Rate (Default: >= 88%)", level=2)

    bold_para(doc, "What it measures: ", "Your monthly attendance rate after adjusting for approved leave.")
    bold_para(doc, "Why it matters: ", "Quality inspection requires consistent presence. Employees who are frequently absent "
              "cannot maintain inspection quality standards.")
    bold_para(doc, "Who it applies to: ", "ALL positions (Type-1 and Type-2). Everyone must meet this condition.")

    doc.add_paragraph()
    bold_para(doc, "How it is calculated:")
    doc.add_paragraph(
        "The system uses the following formula:"
    )
    bullet(doc, "Step 1: Absence Days = Total Working Days - Actual Working Days - Approved Leave Days")
    bullet(doc, "Step 2: Absence Rate = (Absence Days / Total Working Days) x 100")
    bullet(doc, "Step 3: Attendance Rate = 100 - Absence Rate")

    doc.add_paragraph()
    bold_para(doc, "Important rules:")
    bullet(doc, "Approved leave (annual leave, maternity, sick leave AR2, business trip, etc.) does NOT count as absence")
    bullet(doc, "Only AR1 (unauthorized/unapproved absence) counts as absence in this calculation")
    bullet(doc, "Business trips (Di cong tac) count as actual working days")

    add_example_box(doc, "C1 Example — PASS", [
        "Employee: Pham Thi Lan",
        "Month: February 2026 (27 total working days)",
        "Actual working days: 23",
        "Approved leave: 2 days annual leave (Phep nam)",
        "",
        "Absence Days = 27 - 23 - 2 = 2 days",
        "Absence Rate = (2 / 27) x 100 = 7.4%",
        "Attendance Rate = 100 - 7.4 = 92.6%",
        "",
        "92.6% >= 88% --> C1 = PASS",
    ])

    add_example_box(doc, "C1 Example — FAIL", [
        "Employee: Nguyen Van Duc",
        "Month: February 2026 (27 total working days)",
        "Actual working days: 19",
        "Approved leave: 1 day sick leave AR2",
        "",
        "Absence Days = 27 - 19 - 1 = 7 days",
        "Absence Rate = (7 / 27) x 100 = 25.9%",
        "Attendance Rate = 100 - 25.9 = 74.1%",
        "",
        "74.1% < 88% --> C1 = FAIL --> No incentive this month",
    ])

    add_example_box(doc, "C1 Example — Maternity Leave (Approved)", [
        "Employee: Tran Thi Mai",
        "Month: February 2026 (27 total working days)",
        "Actual working days: 15",
        "Approved leave: 10 days maternity leave + 2 days prenatal checkup = 12 days",
        "",
        "Absence Days = 27 - 15 - 12 = 0 days",
        "Absence Rate = (0 / 27) x 100 = 0%",
        "Attendance Rate = 100 - 0 = 100%",
        "",
        "100% >= 88% --> C1 = PASS",
        "All maternity-related leave is approved and does not hurt your attendance.",
    ])

    # ── C2 ──
    doc.add_heading("C2 — Unapproved Absence (Default: <= 2 days)", level=2)

    bold_para(doc, "What it measures: ", "The number of days you were absent WITHOUT approval (AR1 category only).")
    bold_para(doc, "Why it matters: ", "Unapproved absences disrupt team scheduling and production quality.")
    bold_para(doc, "Who it applies to: ", "ALL positions (Type-1 and Type-2).")

    doc.add_paragraph()
    bold_para(doc, "What counts as AR1 (unapproved absence):")
    bullet(doc, "\"Vang khong phep\" — Unauthorized absence (no notice, no approval)")
    bullet(doc, "\"AR1 - Gui thu\" — Written notice absence (letter submitted but not approved leave)")

    bold_para(doc, "What does NOT count as AR1:")
    bullet(doc, "Annual leave, maternity leave, sick leave AR2, business trips, and all other approved leave types")
    bullet(doc, "See the complete leave classification table in Section 5.3.1")

    add_example_box(doc, "C2 Example — PASS", [
        "Employee had 1 day of unauthorized absence (AR1) this month.",
        "1 day <= 2 days threshold --> C2 = PASS",
    ])

    add_example_box(doc, "C2 Example — FAIL", [
        "Employee had 3 days of unauthorized absence (AR1) this month.",
        "3 days > 2 days threshold --> C2 = FAIL --> No incentive",
    ])

    add_example_box(doc, "C2 Example — Annual Leave Does NOT Count", [
        "Employee took 5 days of annual leave (Phep nam) and had 0 AR1 days.",
        "AR1 days = 0 <= 2 days threshold --> C2 = PASS",
        "Annual leave is NOT counted as unapproved absence.",
    ])

    # ── C3 ──
    doc.add_heading("C3 — Actual Working Days (Default: > 0)", level=2)

    bold_para(doc, "What it measures: ", "Whether the employee worked at least 1 day during the month.")
    bold_para(doc, "Why it matters: ", "Employees who were entirely absent for the whole month (e.g., full-month leave) "
              "should not receive incentive.")
    bold_para(doc, "Who it applies to: ", "ALL positions.")

    add_example_box(doc, "C3 Example — PASS", [
        "Employee worked 15 days this month.",
        "15 > 0 --> C3 = PASS",
    ])

    add_example_box(doc, "C3 Example — FAIL", [
        "Employee was on full-month maternity leave. Actual working days = 0.",
        "0 is NOT > 0 --> C3 = FAIL",
        "Note: Even though maternity leave is approved, if you have 0 actual working days,",
        "you are excluded from incentive for that month. This is expected — you were not",
        "performing inspection duties.",
    ])

    # ── C4 ──
    doc.add_heading("C4 — Minimum Working Days (Default: >= 12 days)", level=2)

    bold_para(doc, "What it measures: ", "Whether the employee worked at least 12 days during the month.")
    bold_para(doc, "Why it matters: ", "Employees who work too few days in a month cannot be meaningfully evaluated "
              "on quality performance.")
    bold_para(doc, "Who it applies to: ", "ALL positions.")

    doc.add_paragraph()
    bold_para(doc, "Note: ",
              "This condition ensures that employees with very short working periods (e.g., joined mid-month "
              "or took extended leave) are excluded from incentive for that month.")

    add_example_box(doc, "C4 Example — PASS", [
        "Employee worked 20 days this month.",
        "20 >= 12 --> C4 = PASS",
    ])

    add_example_box(doc, "C4 Example — FAIL", [
        "Employee joined on the 20th of the month and only worked 8 days.",
        "8 < 12 --> C4 = FAIL",
        "This employee will not receive incentive for this month.",
        "Next month, if they work 12+ days, they become eligible.",
    ])

    # ── C5 ──
    doc.add_heading("C5 — Personal AQL Failure (Default: = 0)", level=2)

    bold_para(doc, "What it measures: ", "The number of AQL (Acceptable Quality Level) inspection failures "
              "attributed to you personally in the current month.")
    bold_para(doc, "Why it matters: ", "AQL is the international quality standard used by adidas and other brands. "
              "When shoes fail AQL inspection, it means the entire batch may need to be re-inspected or rejected.")
    bold_para(doc, "Who it applies to: ", "ASSEMBLY INSPECTOR, RQC ASSEMBLY INSPECTOR, AQL INSPECTOR only.")

    doc.add_paragraph()
    bold_para(doc, "How AQL works (simplified): ",
              "A random sample of shoes is taken from a production batch. If the defects in the sample exceed "
              "the acceptable limit, the entire batch FAILS AQL. The inspector who checked that batch gets "
              "an AQL failure record.")

    add_example_box(doc, "C5 Example — PASS", [
        "Inspector Hoa had 0 AQL failures this month.",
        "0 = 0 --> C5 = PASS",
    ])

    add_example_box(doc, "C5 Example — FAIL", [
        "Inspector Duc had 1 AQL failure this month (1 batch failed AQL).",
        "1 != 0 --> C5 = FAIL --> No incentive this month",
        "",
        "Even 1 failure means FAIL. The threshold is ZERO tolerance.",
    ])

    # ── C6 ──
    doc.add_heading("C6 — Personal AQL Consecutive Failures", level=2)

    bold_para(doc, "What it measures: ", "Whether you have had AQL failures in consecutive (back-to-back) months.")
    bold_para(doc, "Why it matters: ", "Repeated failures over multiple months indicate a persistent quality problem "
              "that needs attention.")
    bold_para(doc, "Who it applies to: ", "ASSEMBLY INSPECTOR, RQC ASSEMBLY INSPECTOR only.")

    doc.add_paragraph()
    bold_para(doc, "Year-dependent thresholds:")
    bullet(doc, "For year 2025: 3 or more consecutive months of AQL failure = FAIL")
    bullet(doc, "For year 2026 onwards: 2 or more consecutive months of AQL failure = FAIL")

    doc.add_paragraph(
        "The system looks at the past 3 months of AQL history to determine this condition."
    )

    add_example_box(doc, "C6 Example — PASS (2026 rules)", [
        "Inspector had AQL failure in January 2026 but NOT in February 2026.",
        "Only 1 consecutive month of failure --> C6 = PASS",
    ])

    add_example_box(doc, "C6 Example — FAIL (2026 rules)", [
        "Inspector had AQL failure in January 2026 AND February 2026.",
        "2 consecutive months of failure >= 2 month threshold --> C6 = FAIL",
        "",
        "System tags this as YES_2MONTHS_DEC_JAN or similar consecutive tag.",
    ])

    add_example_box(doc, "C6 Example — PASS (2025 rules, stricter was 3 months)", [
        "Inspector had AQL failure in October 2025 AND November 2025 (but NOT September).",
        "2 consecutive months of failure < 3 month threshold (2025 rules) --> C6 = PASS",
        "Note: Under 2026 rules, this would be FAIL (threshold lowered to 2).",
    ])

    # ── C7 ──
    doc.add_heading("C7 — Team/Area AQL Consecutive Failures", level=2)

    bold_para(doc, "What it measures: ", "Whether employees under your responsibility have had consecutive AQL failures.")
    bold_para(doc, "Why it matters: ", "Line Leaders and Trainers are responsible for their team's quality. "
              "If their team members have persistent quality problems, the leader is also held accountable.")
    bold_para(doc, "Who it applies to: ", "LINE LEADER and AUDITOR & TRAINER only.")
    bold_para(doc, "Exempt: ", "RQC Assembly Inspector (A1B) is exempt from C7 from February 2026.")

    doc.add_paragraph()
    bold_para(doc, "How it works:")
    bullet(doc, "For LINE LEADER: The system checks if any subordinate inspector has consecutive AQL failures")
    bullet(doc, "For AUDITOR & TRAINER: The system checks if any employee in their assigned building area has consecutive AQL failures")
    bullet(doc, "Same year-dependent threshold as C6 (2025: 3+ months, 2026+: 2+ months)")

    add_example_box(doc, "C7 Example — FAIL for Line Leader", [
        "Line Leader Minh's team member, Inspector Hoa, failed AQL in both January and February 2026.",
        "Hoa has 2 consecutive months of failure (2026 threshold: 2 months)",
        "Therefore Minh's C7 = FAIL --> Minh does not receive incentive this month.",
        "",
        "Even though Minh himself did nothing wrong, he is responsible for his team.",
    ])

    # ── C8 ──
    doc.add_heading("C8 — Area Reject Rate (Default: < 3.0%)", level=2)

    bold_para(doc, "What it measures: ", "The AQL reject rate for your assigned area or building.")
    bold_para(doc, "Why it matters: ", "Trainers and Model Masters are responsible for overall area quality, "
              "not just individual inspections.")
    bold_para(doc, "Who it applies to: ", "AUDITOR & TRAINER and MODEL MASTER only.")

    doc.add_paragraph()
    bold_para(doc, "How it works:")
    bullet(doc, "For AUDITOR & TRAINER: The reject rate is calculated for their assigned building area (A, B, C, D, or Repacking)")
    bullet(doc, "For MODEL MASTER: The reject rate is calculated across ALL factories (all-factory reject rate)")
    bullet(doc, "If the reject rate is >= 3.0% (default threshold), C8 = FAIL")

    add_example_box(doc, "C8 Example — PASS", [
        "Trainer Linh is responsible for Building A.",
        "Building A's AQL reject rate this month: 1.8%",
        "1.8% < 3.0% --> C8 = PASS",
    ])

    add_example_box(doc, "C8 Example — FAIL", [
        "Model Master Hoa's all-factory reject rate: 3.5%",
        "3.5% >= 3.0% --> C8 = FAIL --> No incentive for Hoa this month",
    ])

    # ── C9 ──
    doc.add_heading("C9 — 5PRS Pass Rate (Default: >= 95%)", level=2)

    bold_para(doc, "What it measures: ", "Your 5PRS (5 Pairs Random Sampling) inspection pass rate.")
    bold_para(doc, "Why it matters: ", "5PRS is an internal quality check where 5 random pairs of shoes are "
              "inspected from each production order. A high pass rate means you are consistently catching defects.")
    bold_para(doc, "Who it applies to: ", "ASSEMBLY INSPECTOR and RQC ASSEMBLY INSPECTOR only.")

    doc.add_paragraph()
    bold_para(doc, "Formula: ", "5PRS Pass Rate = (Pass Quantity / Total Validation Quantity) x 100")

    add_example_box(doc, "C9 Example — PASS", [
        "Inspector checked 120 pairs total in 5PRS inspections this month.",
        "116 pairs passed, 4 pairs failed.",
        "Pass Rate = (116 / 120) x 100 = 96.7%",
        "96.7% >= 95% --> C9 = PASS",
    ])

    add_example_box(doc, "C9 Example — FAIL", [
        "Inspector checked 100 pairs total in 5PRS inspections this month.",
        "93 pairs passed, 7 pairs failed.",
        "Pass Rate = (93 / 100) x 100 = 93.0%",
        "93.0% < 95% --> C9 = FAIL --> No incentive this month",
    ])

    # ── C10 ──
    doc.add_heading("C10 — 5PRS Inspection Quantity (Default: >= 100 pairs)", level=2)

    bold_para(doc, "What it measures: ", "The total number of pairs you inspected through 5PRS during the month.")
    bold_para(doc, "Why it matters: ", "This ensures inspectors are performing a meaningful volume of inspections, "
              "not just a few token checks.")
    bold_para(doc, "Who it applies to: ", "ASSEMBLY INSPECTOR only. RQC Assembly Inspector (A1B) is EXEMPT from February 2026.")

    add_example_box(doc, "C10 Example — PASS", [
        "Inspector checked 135 pairs this month through 5PRS.",
        "135 >= 100 --> C10 = PASS",
    ])

    add_example_box(doc, "C10 Example — FAIL", [
        "Inspector checked only 72 pairs this month (was on leave for part of the month).",
        "72 < 100 --> C10 = FAIL --> No incentive this month",
        "",
        "Note: If the low quantity was due to area rotation or R&R update,",
        "the manager may use the Allowance system to override this condition.",
    ])

    bold_para(doc, "Why RQC Assembly Inspector (A1B) is exempt: ",
              "RQC inspectors perform process checking and reporting, not line-based inspection. "
              "They do not perform the same volume of 5PRS inspections as regular assembly inspectors, "
              "so the minimum quantity requirement does not apply to them.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # APPENDIX 2
    # ══════════════════════════════════════════
    doc.add_heading("APPENDIX 2: POSITION-CONDITION APPLICATION MATRIX", level=1)

    doc.add_paragraph(
        "This appendix provides the complete position-condition matrix used by the V10 calculation engine. "
        "All positions are configured through JSON configuration files, ensuring consistency across the Python "
        "calculation engine, web dashboard, and validation system."
    )

    doc.add_heading("Type-1 Positions", level=2)
    doc.add_paragraph('Each Type-1 position has a specific set of applicable conditions. "Yes" means the condition must be passed; "-" means not applicable.')

    tbl = doc.add_table(rows=1, cols=11)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Position", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9", "C10"])

    full_matrix = [
        ["ASSEMBLY INSPECTOR", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "Yes"],
        ["RQC ASSEMBLY INSPECTOR\n(A1B, from 2026.02)", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "-"],
        ["AQL INSPECTOR", "Yes", "Yes", "Yes", "Yes", "Yes", "-", "-", "-", "-", "-"],
        ["LINE LEADER", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "-", "-", "-"],
        ["AUDITOR & TRAINER", "Yes", "Yes", "Yes", "Yes", "-", "-", "Yes", "Yes", "-", "-"],
        ["MODEL MASTER", "Yes", "Yes", "Yes", "Yes", "-", "-", "-", "Yes", "-", "-"],
        ["Management\n(GL, Supervisor, Manager)", "Yes", "Yes", "Yes", "Yes", "-", "-", "-", "-", "-", "-"],
    ]
    for m in full_matrix:
        add_row(tbl, m)

    doc.add_heading("Type-2 Positions", level=2)
    doc.add_paragraph(
        "All Type-2 positions apply only attendance conditions (C1-C4). The incentive amount is calculated based on "
        "the average of their corresponding Type-1 position reference, as described in Section 5.7 and 5.8."
    )
    doc.add_paragraph(
        "Type-2 positions include: Line Leader, Assembly Inspector, AQL Inspector, Stitching Inspector, "
        "Bottom Inspector, Cutting Inspector, MTL Inspector, OSC Inspector, OCPT Staff, QA Team, RQC"
    )

    doc.add_heading("Type-3 Positions", level=2)
    doc.add_paragraph(
        "Type-3 (New QIP Member) positions have no applicable conditions and are excluded from incentive payment. "
        "They automatically transition to Type-2 after approximately 3 months."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # APPENDIX 3
    # ══════════════════════════════════════════
    doc.add_heading("APPENDIX 3: HWK QIP TALENT POOL STATUS", level=1)

    doc.add_paragraph(
        "The HWK QIP Talent Pool is managed through the V10 system via the system configuration. "
        "Talent Pool members receive additional monthly incentive independent of their regular Type-1/Type-2 incentive."
    )

    doc.add_paragraph("Current Talent Pool Configuration:")
    bullet(doc, "Auto-apply: Yes — Talent Pool bonus is automatically applied alongside regular incentive.")
    bullet(doc, "Stackable: Yes — Talent Pool bonus stacks with regular incentive.")
    bullet(doc, "Payment timing: With regular incentive cycle.")
    bullet(doc, "Conditions requirement: No — Talent Pool members receive bonus regardless of condition pass/fail.")
    bullet(doc, "Exclusion: Resigned employees are automatically excluded; active status required.")

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: The specific members, amounts, and validity periods are managed by QIP managers and approved by "
        "HWK COO through the standard approval process described in Section 6."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # APPENDIX 4
    # ══════════════════════════════════════════
    doc.add_heading("APPENDIX 4: QIP TEAM R&R MAP", level=1)
    doc.add_paragraph(
        "Please refer to the separate QIP Team R&R Map document for the complete 69-job list mapping of all "
        "QIP team member roles and responsibilities."
    )

    # ══════════════════════════════════════════
    # APPENDIX 5
    # ══════════════════════════════════════════
    doc.add_heading("APPENDIX 5: FIRESTORE DATA COLLECTIONS", level=1)
    doc.add_paragraph("The V10 system uses the following Firestore collections:")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Collection", "Purpose", "Access"])

    collections = [
        ["employees/{monthYear}/all_data/data", "Employee + incentive data", "Auth: read, Admin: write"],
        ["dashboard_summary/{monthYear}", "KPI summary statistics", "Auth: read, Admin: write"],
        ["thresholds/{monthYear}", "7 configurable thresholds", "Auth: read, Admin: write"],
        ["threshold_history/{autoId}", "Immutable audit trail", "Auth: read, Admin: create"],
        ["allowances/{monthYear}/items/{docId}", "Exception overrides", "Auth: read, Admin: write"],
        ["system/config", "System settings", "Auth: read, Admins: write"],
        ["configs/{document}", "Position mapping, config", "Auth: read, Admin: write"],
        ["system_feedback/{autoId}", "Feedback / issue tracking", "Auth: read/create, Admin: full"],
        ["pendingNotifications/{autoId}", "Email notification queue", "Auth: create, Admin: read"],
        ["config/email", "SMTP credentials", "Admin only"],
        ["email_logs/{autoId}", "Email delivery audit", "Admin only"],
    ]
    for c in collections:
        add_row(tbl, c)

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- The End ---")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    return doc


if __name__ == "__main__":
    doc = create_policy()
    output_path = os.path.expanduser(
        "~/Downloads/2025 QIP INCENTIVE POLICY Version3_Mar2026-Feb2027_En.docx"
    )
    doc.save(output_path)
    print(f"Policy document saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path):,} bytes")
