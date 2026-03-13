#!/usr/bin/env python3
"""
Generate QIP Incentive Policy Document Version 3
- 100% aligned with HWK QIP Incentive System V10
- Validity: March 2026 ~ February 2027
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table_row(table, cells_data, bold=False, header=False, bg_color=None):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
        if bg_color:
            set_cell_shading(cell, bg_color)
        elif header:
            set_cell_shading(cell, "2F5496")
    return row

def create_policy():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ══════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HWK VINA")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("QIP INCENTIVE POLICY")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run("Version 3.0")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    validity = doc.add_paragraph()
    validity.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = validity.add_run("Validity Period: March 2026 ~ February 2027")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Updated: March 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dept.add_run("Quality Improvement Program Department")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # DOCUMENT CONTROL
    # ══════════════════════════════════════════
    doc.add_heading("Document Control", level=1)

    # Version history table
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Version", "Date", "Changes", "Author"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    versions = [
        ["1.0", "Mar 2025", "Initial QIP Incentive Policy", "QIP Department"],
        ["2.0", "Mar 2026", "Updated to reflect V10 system, added system architecture, allowance system, position-condition matrix", "QIP Department"],
        ["3.0", "Mar 2026", "Full alignment with V10 system code. Added: RQC A1B exemptions, LINE LEADER 12% formula, feedback system details, email notifications, admin panel features, CI/CD pipeline, consecutive failure year-dependent thresholds", "QIP Department"],
    ]
    for v in versions:
        add_table_row(tbl, v)

    doc.add_paragraph()

    # Approval table
    doc.add_heading("Approval", level=2)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = 'Table Grid'
    hdr2 = tbl2.rows[0].cells
    for i, text in enumerate(["Role", "Name", "Signature", "Date"]):
        p = hdr2[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr2[i], "2F5496")

    approvals = [
        ["Prepared by", "", "", ""],
        ["Reviewed by", "", "", ""],
        ["Approved by", "", "", ""],
    ]
    for a in approvals:
        add_table_row(tbl2, a)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Purpose & Scope",
        "2. Core Principle",
        "3. Employee Classification (TYPE System)",
        "4. 10 Incentive Conditions (C1–C10)",
        "5. Configurable Thresholds",
        "6. Position-Condition Matrix",
        "7. Incentive Calculation",
        "   7.1 TYPE-1: Progressive Table",
        "   7.2 TYPE-1: LINE LEADER Calculation",
        "   7.3 TYPE-1: AQL Inspector 3-Part Calculation",
        "   7.4 TYPE-2: Multiplier-Based Calculation",
        "   7.5 TYPE-3: New Members (Excluded)",
        "   7.6 Talent Pool Bonus (Program 2)",
        "8. Consecutive Failure Policy",
        "9. Allowance / Exception Override System",
        "10. System Architecture & Dashboard (V10)",
        "   10.1 Data Flow & CI/CD Pipeline",
        "   10.2 Dashboard Features (8 Tabs)",
        "   10.3 Admin Panel (5 Tabs)",
        "   10.4 Threshold Management",
        "   10.5 Feedback System",
        "   10.6 Email Notification System",
        "   10.7 Multi-Language Support",
        "   10.8 Authentication & RBAC",
        "11. Appendix",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10)
        if not item.startswith("   "):
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 1. PURPOSE & SCOPE
    # ══════════════════════════════════════════
    doc.add_heading("1. Purpose & Scope", level=1)

    doc.add_paragraph(
        "This document defines the QIP (Quality Improvement Program) Incentive Policy for HWK Vina. "
        "It describes the conditions, calculations, and systems used to evaluate and distribute monthly "
        "incentives to eligible employees based on quality, attendance, and performance criteria."
    )
    doc.add_paragraph(
        "The policy covers all QIP-enrolled positions including production inspectors, line leaders, "
        "managers, auditors, and support staff. The incentive system is designed to reward consistent "
        "quality performance and encourage continuous improvement."
    )

    p = doc.add_paragraph()
    run = p.add_run("Validity Period: ")
    run.bold = True
    run = p.add_run("March 2026 to February 2027")

    p = doc.add_paragraph()
    run = p.add_run("System Version: ")
    run.bold = True
    run = p.add_run("V10 (Firestore-based, GitHub Pages deployment)")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 2. CORE PRINCIPLE
    # ══════════════════════════════════════════
    doc.add_heading("2. Core Principle", level=1)

    p = doc.add_paragraph()
    run = p.add_run("100% Condition Fulfillment Required")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph(
        "An employee must pass ALL applicable conditions (100% pass rate) to receive the incentive. "
        "Passing 80% or 99% of conditions does NOT qualify. Only 100% compliance earns the incentive."
    )
    doc.add_paragraph(
        "The only exception is when a manager applies an Allowance (Exception Override) through the "
        "admin panel, which can override specific failed conditions with documented justification."
    )

    # ══════════════════════════════════════════
    # 3. EMPLOYEE CLASSIFICATION
    # ══════════════════════════════════════════
    doc.add_heading("3. Employee Classification (TYPE System)", level=1)

    doc.add_paragraph(
        "All QIP-enrolled employees are classified into one of three types, which determines their "
        "applicable conditions and calculation method:"
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["TYPE", "Description", "Applicable Conditions", "Calculation Method"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    types = [
        ["TYPE-1", "Direct quality workers\n(Inspectors, Line Leaders, Auditors, Model Masters)", "C1–C10\n(varies by position)", "Progressive Table or\nSubordinate-based"],
        ["TYPE-2", "Management & support\n(Group Leader, Supervisor, Manager, etc.)", "C1–C4 only\n(attendance)", "Multiplier × LINE LEADER average"],
        ["TYPE-3", "New members\n(< 3 months tenure)", "None", "0 VND (excluded)\nUpgrades to TYPE-2 after 3 months"],
    ]
    for t in types:
        add_table_row(tbl, t)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 4. 10 INCENTIVE CONDITIONS
    # ══════════════════════════════════════════
    doc.add_heading("4. 10 Incentive Conditions (C1–C10)", level=1)

    doc.add_paragraph(
        "The system evaluates up to 10 conditions per employee. Which conditions apply depends on "
        "the employee's position (see Section 6: Position-Condition Matrix)."
    )

    # Table 1: Conditions
    doc.add_heading("Table 1: Condition Definitions", level=2)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["#", "Condition", "Category", "Default Threshold", "Description"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    conditions = [
        ["C1", "Attendance Rate", "Attendance", "≥ 88%", "Attendance rate with approved-leave adjustment. Calculated as: (Actual Working Days) / (Total Working Days − Approved Leave Days) × 100%"],
        ["C2", "Unapproved Absence", "Attendance", "≤ 2 days", "AR1 (unapproved) absence days must be within limit"],
        ["C3", "Actual Working Days", "Attendance", "> 0 days", "Must have at least 1 actual working day in the month"],
        ["C4", "Minimum Working Days", "Attendance", "≥ 12 days", "Only enforced after 20th of current month. Interim reports (before 20th) exempt this condition"],
        ["C5", "Personal AQL Failure", "AQL", "= 0 failures", "No personal AQL inspection failures in the current month"],
        ["C6", "Personal AQL Consecutive", "AQL", "No consecutive failures", "No consecutive month AQL failures. Threshold: 2025 = 3+ months, 2026+ = 2+ months"],
        ["C7", "Team/Area AQL Consecutive", "AQL", "No consecutive failures", "No consecutive month team/area AQL failures. Same year-dependent threshold as C6"],
        ["C8", "Area Reject Rate", "AQL", "< 3.0%", "Reject rate for assigned area/building must be below threshold"],
        ["C9", "5PRS Pass Rate", "5PRS", "≥ 95%", "5PRS quality inspection pass rate must meet threshold"],
        ["C10", "5PRS Inspection Quantity", "5PRS", "≥ 100 pairs", "Minimum 5PRS inspection quantity required"],
    ]
    for c in conditions:
        add_table_row(tbl, c)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "All thresholds shown above are system defaults. Administrators can adjust thresholds "
        "monthly via the Admin Panel. See Section 5 for configurable threshold details."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 5. CONFIGURABLE THRESHOLDS
    # ══════════════════════════════════════════
    doc.add_heading("5. Configurable Thresholds", level=1)

    doc.add_paragraph(
        "Seven thresholds can be adjusted monthly by administrators through the Admin Panel. "
        "Changes are stored in Firestore with an immutable audit trail."
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Threshold", "System Default", "Condition", "Field Key"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    thresholds = [
        ["Attendance Rate", "88%", "C1", "attendance_rate"],
        ["Unapproved Absence", "2 days", "C2", "unapproved_absence"],
        ["Minimum Working Days", "12 days", "C4", "minimum_working_days"],
        ["Area Reject Rate", "3.0%", "C8", "area_reject_rate"],
        ["5PRS Pass Rate", "95%", "C9", "5prs_pass_rate"],
        ["5PRS Min Quantity", "100 pairs", "C10", "5prs_min_qty"],
        ["Consecutive AQL Months", "3 months", "C6, C7", "consecutive_aql_months"],
    ]
    for t in thresholds:
        add_table_row(tbl, t)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Audit Trail: ")
    run.bold = True
    p.add_run(
        "Every threshold change is recorded in the threshold_history collection (immutable). "
        "Records include: changed fields, old values, new values, changed by (email), and timestamp."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 6. POSITION-CONDITION MATRIX
    # ══════════════════════════════════════════
    doc.add_heading("6. Position-Condition Matrix", level=1)

    doc.add_paragraph(
        "Not all conditions apply to every position. The matrix below shows which conditions "
        "are evaluated for each position category."
    )

    # Table: TYPE-1 Matrix
    doc.add_heading("Table 2: TYPE-1 Position-Condition Matrix", level=2)

    tbl = doc.add_table(rows=1, cols=11)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    headers = ["Position", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9", "C10"]
    for i, text in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    matrix = [
        ["ASSEMBLY INSPECTOR", "Y", "Y", "Y", "Y", "Y", "Y", "—", "—", "Y", "Y"],
        ["RQC ASSEMBLY INSPECTOR\n(A1B, from 2026.02)", "Y", "Y", "Y", "Y", "Y", "Y", "—", "—", "Y", "—"],
        ["AQL INSPECTOR", "Y", "Y", "Y", "Y", "Y", "—", "—", "—", "—", "—"],
        ["AUDITOR & TRAINER", "Y", "Y", "Y", "Y", "—", "—", "Y", "Y", "—", "—"],
        ["MODEL MASTER", "Y", "Y", "Y", "Y", "—", "—", "—", "Y", "—", "—"],
        ["LINE LEADER", "Y", "Y", "Y", "Y", "—", "—", "Y", "—", "—", "—"],
        ["Management\n(Manager, Supervisor,\nGroup Leader, etc.)", "Y", "Y", "Y", "Y", "—", "—", "—", "—", "—", "—"],
    ]
    for m in matrix:
        row = add_table_row(tbl, m)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Y")
    run.bold = True
    p.add_run(" = Condition applies,  ")
    run = p.add_run("—")
    run.bold = True
    p.add_run(" = Condition not applicable (automatically passed)")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("RQC Assembly Inspector (A1B): ")
    run.bold = True
    p.add_run(
        "From February 2026 onwards, position code A1B is classified as RQC_ASSEMBLY_INSPECTOR. "
        "C7 (Team/Area AQL) and C10 (5PRS Inspection Qty) are exempted because RQC inspectors "
        "perform process checking and reports, not line inspection. Before February 2026, A1B is "
        "treated as regular ASSEMBLY INSPECTOR."
    )

    # Table: TYPE-2
    doc.add_heading("Table 3: TYPE-2 Position-Condition Matrix", level=2)
    doc.add_paragraph(
        "All TYPE-2 positions are evaluated on attendance conditions only (C1–C4). "
        "AQL and 5PRS conditions (C5–C10) do not apply."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 7. INCENTIVE CALCULATION
    # ══════════════════════════════════════════
    doc.add_heading("7. Incentive Calculation", level=1)

    # 7.1 Progressive Table
    doc.add_heading("7.1 TYPE-1: Progressive Table", level=2)

    doc.add_paragraph(
        "TYPE-1 employees who pass all applicable conditions receive incentive based on "
        "consecutive months of qualification. The amount increases progressively:"
    )

    doc.add_heading("Table 4: Progressive Incentive Table", level=3)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Consecutive Months", "Amount (VND)"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    progressive = [
        ["1", "150,000"],
        ["2", "250,000"],
        ["3", "300,000"],
        ["4", "350,000"],
        ["5", "400,000"],
        ["6", "450,000"],
        ["7", "500,000"],
        ["8", "650,000"],
        ["9", "750,000"],
        ["10", "850,000"],
        ["11", "950,000"],
        ["12–15", "1,000,000 (maximum)"],
    ]
    for p_row in progressive:
        add_table_row(tbl, p_row)

    doc.add_paragraph()
    doc.add_paragraph("Applies to: ASSEMBLY INSPECTOR, AUDITOR & TRAINER, MODEL MASTER")
    p = doc.add_paragraph()
    run = p.add_run("Reset rule: ")
    run.bold = True
    p.add_run("If an employee fails any condition in a month, consecutive months resets to 0.")

    # 7.2 LINE LEADER
    doc.add_heading("7.2 TYPE-1: LINE LEADER Calculation", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Formula: ")
    run.bold = True

    doc.add_paragraph(
        "LINE LEADER Incentive = Subordinate Incentive Total × 12%"
    )
    doc.add_paragraph(
        "The system uses a subordinate mapping to identify all team members under each LINE LEADER. "
        "Only subordinates who received incentive (> 0 VND) are counted in the total. "
        "The LINE LEADER must also pass all applicable conditions (C1, C2, C3, C4, C7) to receive the incentive."
    )

    # 7.3 AQL Inspector 3-Part
    doc.add_heading("7.3 TYPE-1: AQL Inspector 3-Part Calculation", level=2)

    doc.add_paragraph(
        "AQL Inspectors receive a special 3-part incentive structure. All three parts are tracked separately."
    )

    p = doc.add_paragraph()
    run = p.add_run("Total AQL Incentive = Part 1 + Part 2 + Part 3")
    run.bold = True

    doc.add_heading("Part 1: AQL Inspection Evaluation", level=3)
    doc.add_paragraph(
        "Based on the same Progressive Table (Table 4). Consecutive months of sustained "
        "performance (rejection rate < 3% by adidas/T1QM/3rd Party) earns the progressive amount."
    )

    doc.add_heading("Part 2: CFA Certificate Incentive", level=3)
    p = doc.add_paragraph()
    p.add_run("Condition: Valid adidas-approved AQL certificate (CFA certified)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Fixed Amount: 700,000 VND per month")
    run.bold = True

    doc.add_heading("Part 3: HWK Claim Prevention Incentive", level=3)
    doc.add_paragraph("Condition: No HWK complaints. Progressive schedule:")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Consecutive Months", "Amount (VND)"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    hwk_table = [
        ["0–3", "0"],
        ["4–6", "300,000"],
        ["7–9", "500,000"],
        ["10–12", "700,000"],
        ["13–15", "900,000"],
    ]
    for h in hwk_table:
        add_table_row(tbl, h)

    doc.add_paragraph()

    # 7.4 TYPE-2
    doc.add_heading("7.4 TYPE-2: Multiplier-Based Calculation", level=2)

    doc.add_paragraph(
        "TYPE-2 employees receive incentive calculated as a multiplier of the TYPE-1 LINE LEADER average. "
        "The LINE LEADER average is calculated from receiving members only (employees with incentive > 0; zeros excluded)."
    )

    doc.add_heading("Table 5: TYPE-2 Multipliers", level=3)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Position", "Multiplier", "Formula"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    multipliers = [
        ["GROUP LEADER", "2.0×", "LINE LEADER receiving average × 2.0"],
        ["SUPERVISOR /\n(V) SUPERVISOR /\nA.SUPERVISOR", "2.5×", "LINE LEADER receiving average × 2.5"],
        ["A.MANAGER\n(Assistant Manager)", "3.0×", "LINE LEADER receiving average × 3.0"],
        ["MANAGER", "3.5×", "LINE LEADER receiving average × 3.5"],
        ["S.MANAGER\n(Senior Manager)", "4.0×", "LINE LEADER receiving average × 4.0"],
    ]
    for m in multipliers:
        add_table_row(tbl, m)

    doc.add_paragraph()

    doc.add_heading("TYPE-2 Non-Management Position Mapping", level=3)
    doc.add_paragraph(
        "Some TYPE-2 positions use a different base average for their calculation:"
    )

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Position", "Base Average Used"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    mappings = [
        ["BOTTOM / CUTTING / MTL / STITCHING /\nOSC INSPECTOR, OCPT STAFF, RQC", "ASSEMBLY INSPECTOR average"],
        ["AQL INSPECTOR (TYPE-2)", "ASSEMBLY INSPECTOR average"],
        ["QA TEAM (default)", "ASSEMBLY INSPECTOR average"],
        ["QA TEAM (code QA3A)", "GROUP LEADER average"],
        ["QA TEAM (code QA3B)", "ASSEMBLY INSPECTOR average"],
        ["LINE LEADER (TYPE-2)", "TYPE-1 LINE LEADER average"],
        ["(V) SUPERVISOR (TYPE-2)", "TYPE-1 (VICE) SUPERVISOR average"],
    ]
    for m in mappings:
        add_table_row(tbl, m)

    doc.add_page_break()

    # 7.5 TYPE-3
    doc.add_heading("7.5 TYPE-3: New Members (Excluded)", level=2)
    doc.add_paragraph(
        "New employees with less than 3 months of tenure are classified as TYPE-3. "
        "They are excluded from incentive calculations (0 VND). No conditions are evaluated. "
        "After 3 months, they are automatically upgraded to TYPE-2."
    )

    # 7.6 Talent Pool
    doc.add_heading("7.6 Talent Pool Bonus (Program 2)", level=2)
    doc.add_paragraph(
        "The Talent Pool bonus is a supplementary program for selected employees:"
    )
    doc.add_paragraph("• Bonus is added ON TOP of regular incentive (stackable)", style='List Bullet')
    doc.add_paragraph("• No condition requirements (bonus paid regardless of condition pass/fail)", style='List Bullet')
    doc.add_paragraph("• Members are manually registered in the system configuration", style='List Bullet')
    doc.add_paragraph("• Each member has a defined bonus amount and validity period", style='List Bullet')
    doc.add_paragraph("• Resigned employees are automatically excluded", style='List Bullet')
    doc.add_paragraph("• Payment timing: Together with regular monthly incentive", style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 8. CONSECUTIVE FAILURE POLICY
    # ══════════════════════════════════════════
    doc.add_heading("8. Consecutive Failure Policy", level=1)

    doc.add_paragraph(
        "Conditions C6 and C7 evaluate consecutive AQL failures. The threshold depends on the year:"
    )

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Year", "Threshold", "Rule"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    years = [
        ["2025", "3+ consecutive months", "Only 3-month or longer consecutive failures cause condition fail"],
        ["2026 onwards", "2+ consecutive months", "2-month or longer consecutive failures cause condition fail (stricter)"],
    ]
    for y in years:
        add_table_row(tbl, y)

    doc.add_paragraph()
    doc.add_paragraph(
        "The system automatically applies the correct threshold based on the evaluation year. "
        "This transition was implemented to progressively raise quality standards."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 9. ALLOWANCE / EXCEPTION OVERRIDE
    # ══════════════════════════════════════════
    doc.add_heading("9. Allowance / Exception Override System", level=1)

    doc.add_paragraph(
        "The Allowance system allows administrators to override specific failed conditions "
        "for individual employees with documented justification. This is used for exceptional cases "
        "where condition failure was not due to the employee's own fault."
    )

    doc.add_heading("9.1 Exception Categories", level=2)

    doc.add_paragraph("1. Medical reasons (illness, hospitalization)", style='List Bullet')
    doc.add_paragraph("2. Company order (reassignment, mandatory tasks)", style='List Bullet')
    doc.add_paragraph("3. Natural disaster (flooding, typhoon)", style='List Bullet')
    doc.add_paragraph("4. Other documented reasons (case-by-case)", style='List Bullet')

    doc.add_heading("9.2 Reason Codes in System", level=2)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Code", "Description"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    codes = [
        ["MEDICAL", "Medical reason (illness, injury, hospitalization)"],
        ["COMPANY_ORDER", "Company order (reassignment, mandatory tasks)"],
        ["NATURAL_DISASTER", "Natural disaster (flooding, typhoon, etc.)"],
        ["OTHER", "Other documented reason (free text required)"],
    ]
    for c in codes:
        add_table_row(tbl, c)

    doc.add_heading("9.3 Workflow", level=2)

    steps = [
        "1. Administrator searches for the employee in the Allowance tab",
        "2. System displays current condition status (pass/fail for each applicable condition)",
        "3. Administrator selects the failed conditions to override",
        "4. Administrator provides reason code and detailed justification (required)",
        "5. Preview screen shows the override impact",
        "6. On Apply: system updates employee data, recalculates incentive amount, updates dashboard summary",
        "7. Allowance can be revoked later (restores original condition status)",
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_heading("9.4 Audit Trail", level=2)
    doc.add_paragraph(
        "Every allowance action is recorded with: employee number, employee name, position, "
        "overridden conditions, reason code, reason detail, applied by (admin email), "
        "applied timestamp, and status (APPLIED / REVOKED)."
    )

    doc.add_heading("9.5 Incentive Recalculation", level=2)
    doc.add_paragraph(
        "When an allowance is applied, the system automatically recalculates the incentive amount. "
        "If the employee's continuous_months was 0 (because the pipeline calculated before the allowance), "
        "the system recalculates from previous_continuous_months or reverse-calculates from previous_incentive amount. "
        "A bulk recalculation button is also available to fix all allowance employees with 0 incentive at once."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 10. SYSTEM ARCHITECTURE & DASHBOARD
    # ══════════════════════════════════════════
    doc.add_heading("10. System Architecture & Dashboard (V10)", level=1)

    doc.add_paragraph(
        "The QIP Incentive System V10 is a Firestore-based web application deployed on GitHub Pages. "
        "It replaced V9 (which had security issues with inline employee data in public HTML files) "
        "with a secure architecture where the HTML is a UI shell only and all data is loaded via "
        "authenticated Firestore queries."
    )

    # 10.1 Data Flow
    doc.add_heading("10.1 Data Flow & CI/CD Pipeline", level=2)

    doc.add_paragraph("The automated data pipeline runs every 6 hours (and on manual trigger):")
    doc.add_paragraph()

    steps = [
        "1. Download input data from Google Drive (attendance, AQL results, 5PRS data)",
        "2. Sync thresholds from Firestore (admin-changed values)",
        "3. Sync configuration settings from Firestore",
        "4. Auto-detect stale months requiring recalculation",
        "5. Convert attendance data format",
        "6. Validate attendance data schema",
        "7. Download previous month incentive data",
        "8. Run Python calculation engine (evaluate all 10 conditions for all employees)",
        "9. Upload results to Firestore",
        "10. Process pending email notifications via SMTP",
        "11. Generate selector page and deploy to GitHub Pages",
    ]
    for s in steps:
        doc.add_paragraph(s)

    p = doc.add_paragraph()
    run = p.add_run("Technology Stack: ")
    run.bold = True
    p.add_run(
        "Python 3.11 (calculation engine), Firebase/Firestore (database), "
        "GitHub Actions (CI/CD), GitHub Pages (hosting), "
        "Cloud Functions v2 Node.js 22 (email notifications), "
        "Vanilla JS + Bootstrap 5.3 + Chart.js (frontend)"
    )

    # 10.2 Dashboard
    doc.add_heading("10.2 Dashboard Features (8 Tabs)", level=2)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["#", "Tab", "Features"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    tabs = [
        ["1", "Summary", "KPI cards (recipients, payment rate, total amount), TYPE distribution table, condition pass/fail charts, trend chart, talent pool display"],
        ["2", "Position Detail", "Position-specific breakdown tables with condition status"],
        ["3", "Individual Detail", "Per-employee search with filters, sorting, pagination (250ms debounce). Employee detail modal with full condition breakdown"],
        ["4", "Incentive Criteria", "Reference display of current thresholds and condition definitions"],
        ["5", "Organization Chart", "Hierarchical org chart with expected incentive calculations per position"],
        ["6", "Team Management", "Team-level analytics and performance comparison"],
        ["7", "Validation Summary", "12 KPI validation cards, data integrity checks, system verification"],
        ["8", "Attendance Lookup", "Individual attendance calendar with daily status lookup"],
    ]
    for t in tabs:
        add_table_row(tbl, t)

    # 10.3 Admin Panel
    doc.add_heading("10.3 Admin Panel (5 Tabs)", level=2)

    doc.add_paragraph("The Admin Panel is accessible to administrators only (role-based access control).")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["#", "Tab", "Features"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    admin_tabs = [
        ["1", "Thresholds", "Adjust 7 configurable thresholds per month. View change history with full audit trail."],
        ["2", "Config Management", "Position mapping, progressive table configuration, type classification rules"],
        ["3", "System", "Pipeline status monitoring, system configuration, admin email list management"],
        ["4", "Data Lookup", "Direct Firestore data querying and verification"],
        ["5", "Allowances", "Exception override management with apply/revoke workflow and bulk recalculation"],
    ]
    for t in admin_tabs:
        add_table_row(tbl, t)

    # 10.4 Threshold Management
    doc.add_heading("10.4 Threshold Management", level=2)
    doc.add_paragraph(
        "Thresholds are stored centrally in Firestore (thresholds/{monthYear} collection). "
        "Administrators adjust values through the Admin Dashboard. Every change creates an immutable "
        "record in the threshold_history collection. The Python calculation engine syncs these values "
        "before each run, ensuring calculations always use the latest admin-approved thresholds."
    )

    # 10.5 Feedback System
    doc.add_heading("10.5 Feedback System", level=2)
    doc.add_paragraph(
        "The system includes a built-in feedback/issue tracking system accessible to all authenticated users."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Feedback Types: ")
    run.bold = True
    p.add_run("BUG, IMPROVEMENT, NEW_FEATURE, UI_UX, DATA, OTHER")

    p = doc.add_paragraph()
    run = p.add_run("Priorities: ")
    run.bold = True
    p.add_run("Critical, High, Medium, Low")

    p = doc.add_paragraph()
    run = p.add_run("Status Flow: ")
    run.bold = True
    p.add_run("SUBMITTED → REVIEWING → IN_PROGRESS → COMPLETED (or REJECTED)")

    p = doc.add_paragraph()
    run = p.add_run("Features: ")
    run.bold = True
    p.add_run(
        "Image attachments (up to 3, base64), multiple notification recipients, "
        "admin status management, admin reply with email notification"
    )

    # 10.6 Email Notifications
    doc.add_heading("10.6 Email Notification System", level=2)
    doc.add_paragraph(
        "Three automated email notifications are implemented via Cloud Functions (Node.js 22, asia-northeast3):"
    )

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Trigger", "Recipients", "Content"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    emails = [
        ["New feedback submitted", "All administrators", "Feedback details with type, priority, description (3-language: ko/en/vi)"],
        ["Feedback status changed", "Feedback reporter", "Status update notification with new status"],
        ["Admin replies to feedback", "Feedback reporter", "Reply content with admin comment"],
    ]
    for e in emails:
        add_table_row(tbl, e)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("SMTP: ")
    run.bold = True
    p.add_run("mail.hsvina.com:465 (SSL). All emails logged to email_logs collection for audit.")

    # 10.7 i18n
    doc.add_heading("10.7 Multi-Language Support", level=2)
    doc.add_paragraph(
        "The entire dashboard supports three languages: Korean (ko), English (en), and Vietnamese (vi). "
        "All UI strings, error messages, and labels are managed through the i18n system. "
        "Users can switch language at any time via the language selector."
    )

    # 10.8 Auth
    doc.add_heading("10.8 Authentication & RBAC", level=2)
    doc.add_paragraph(
        "Authentication uses Firebase Email/Password authentication. "
        "Role-Based Access Control (RBAC) is implemented with two levels:"
    )
    doc.add_paragraph("• Regular users: Read-only dashboard access + feedback submission", style='List Bullet')
    doc.add_paragraph("• Administrators: Full access including admin panel, threshold management, allowance overrides, and feedback management", style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph(
        "Admin email list is stored dynamically in Firestore (system/config.admin_emails) "
        "and can be managed through the System tab in the Admin Panel."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 11. APPENDIX
    # ══════════════════════════════════════════
    doc.add_heading("11. Appendix", level=1)

    doc.add_heading("A. Firestore Collections", level=2)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Collection", "Purpose", "Access"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], "2F5496")

    collections = [
        ["employees/{monthYear}/all_data/data", "Employee + incentive data (~270KB)", "Auth: read, Admin: write"],
        ["dashboard_summary/{monthYear}", "KPI summary statistics", "Auth: read, Admin: write"],
        ["thresholds/{monthYear}", "7 configurable thresholds", "Auth: read, Admin: write"],
        ["threshold_history/{autoId}", "Immutable audit trail", "Auth: read, Admin: create only"],
        ["allowances/{monthYear}/items/{docId}", "Exception overrides", "Auth: read, Admin: write"],
        ["system/config", "System settings (admin emails)", "Auth: read, Specific admins: write"],
        ["configs/{document}", "Configuration (position mapping)", "Auth: read, Admin: write"],
        ["system_feedback/{autoId}", "Feedback/issue tracking", "Auth: read/create, Admin: full"],
        ["pendingNotifications/{autoId}", "Email notification queue", "Auth: create, Admin: read/update"],
        ["config/email", "SMTP credentials", "Admin only"],
        ["email_logs/{autoId}", "Email delivery audit", "Admin only"],
    ]
    for c in collections:
        add_table_row(tbl, c)

    doc.add_paragraph()

    doc.add_heading("B. Auditor/Trainer Area Mapping", level=2)
    doc.add_paragraph(
        "Each Auditor & Trainer is assigned to specific building areas for C8 (Area Reject Rate) "
        "evaluation. The mapping is maintained in the system configuration. "
        "Model Masters are evaluated against ALL areas (entire factory). "
        "Default reject rate threshold: configurable per month via admin panel."
    )

    doc.add_heading("C. Special Edge Cases", level=2)
    doc.add_paragraph("• Interim reports (before 20th): C4 (minimum 12 days) is automatically exempted", style='List Bullet')
    doc.add_paragraph("• Business trips (Di cong tac): Counted as actual working days", style='List Bullet')
    doc.add_paragraph("• Stop Working employees: Processed with 0 attendance (not excluded from dataset)", style='List Bullet')
    doc.add_paragraph("• QA3A position code: Treated as GROUP LEADER for TYPE-2 calculation", style='List Bullet')
    doc.add_paragraph("• Approved leave: Deducted from total working days before attendance rate calculation", style='List Bullet')

    doc.add_heading("D. Calculation Execution Order", level=2)
    doc.add_paragraph("The Python engine processes incentives in this order:")
    order = [
        "1. Data validation and previous month data loading",
        "2. Approved leave adjustment + condition evaluation for all employees",
        "3. TYPE-1 calculations (Assembly Inspector, Auditor/Trainer, Model Master, AQL Inspector)",
        "4. TYPE-1 LINE LEADER calculation (subordinate-based)",
        "5. TYPE-2 GROUP LEADER calculation",
        "6. TYPE-2 MANAGER calculation (Supervisor, A.Manager, Manager, S.Manager)",
        "7. TYPE-2 non-management position calculation",
        "8. TYPE-3 new member assignment (0 VND)",
        "9. Talent Pool bonus application",
        "10. Dashboard summary generation",
    ]
    for o in order:
        doc.add_paragraph(o)

    # ── Footer note ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("QIP Incentive Policy Version 3.0 | HWK Vina | March 2026 – February 2027")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == "__main__":
    doc = create_policy()
    output_path = os.path.expanduser(
        "~/Downloads/2025 QIP INCENTIVE POLICY Version3_Mar2026-Feb2027_En.docx"
    )
    doc.save(output_path)
    print(f"Policy document saved to: {output_path}")
