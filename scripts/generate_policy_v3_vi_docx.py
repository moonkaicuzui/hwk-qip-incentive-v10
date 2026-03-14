#!/usr/bin/env python3
"""
Generate QIP Incentive Policy Document Version 3 — Vietnamese
Massively expanded version with detailed condition explanations,
calculation examples, screenshots, and simple language for factory workers.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
SCREENSHOT_DIR = os.path.join(PROJECT_DIR, 'policy_screenshots')


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_header_row(table, headers, bg="2F5496"):
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], bg)


def add_row(table, cells_data, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold:
            run.bold = True
        if bg_color:
            set_cell_shading(cell, bg_color)
    return row


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def bold_para(doc, bold_text, normal_text=""):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_screenshot(doc, filename, caption, width=Inches(6)):
    """Add a screenshot image with caption. Skip gracefully if file not found."""
    img_path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[Hình ảnh: {caption} — Không tìm thấy file {filename}]")
        run.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)


def add_example_box(doc, title, lines):
    """Add a highlighted example box using a single-cell table with background color."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "E8F0FE")  # Light blue background

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(47, 84, 150)

    # Content lines
    for line in lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after box


def create_policy_vi():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ══════════════════════════════════════════
    # TRANG BIA
    # ══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CHINH SACH THUONG QIP CUA HWK")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phien ban 3 — Cap nhat Thang 3/2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thoi han hieu luc: Thang 3/2026 ~ Thang 2/2027")
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Huong dan day du cho nhan vien QIP\n"
        "Bo phan Chuong trinh Cai tien Chat luong (QIP)\n"
        "HWK Vina"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # MUC LUC
    # ══════════════════════════════════════════
    doc.add_heading("MUC LUC", level=1)

    toc_items = [
        ("Phan 1.", " Tong quan — Thuong QIP la gi?"),
        ("Phan 2.", " Giai thich chung ve cau truc chinh sach thuong"),
        ("Phan 3.", " Pham vi va doi tuong ap dung"),
        ("Phan 4.", " Dinh nghia & Thuat ngu viet tat"),
        ("Phan 5.", " Phuong phap tinh thuong — Chuong trinh 1 (co vi du tinh toan)"),
        ("Phan 6.", " Phuong phap tinh thuong — Chuong trinh 2"),
        ("Phan 7.", " Tom tat cac dieu khoan ngoai le"),
        ("Phan 8.", " He thong & Bang dieu khien (V10) — co hinh anh minh hoa"),
        ("Phu luc 1:", " 10 Dieu kien thuong — Giai thich CHI TIET voi vi du"),
        ("Phu luc 2:", " Ma tran ap dung dieu kien theo chuc vu"),
        ("Phu luc 3:", " Trang thai nhom Talent Pool QIP HWK"),
        ("Phu luc 4:", " Ban do R&R doi QIP"),
        ("Phu luc 5:", " Bo suu tap du lieu Firestore"),
    ]
    for bold_part, normal_part in toc_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(normal_part)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 1. TONG QUAN
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 1. TONG QUAN — THUONG QIP LA GI?", level=1)

    doc.add_paragraph(
        'Chao mung ban den voi chuong trinh thuong QIP cua HWK! '
        'Tai lieu nay se giai thich cho ban hieu ve chuong trinh thuong chat luong — '
        'mot chinh sach dac biet de thuong cho nhung nhan vien lam tot cong viec kiem tra chat luong.'
    )

    doc.add_paragraph(
        'Noi don gian: Neu ban di lam day du, khong co loi chat luong, va lam dung cong viec — '
        'ban se duoc nhan them tien thuong moi thang. Cang lam tot lau, tien thuong cang tang!'
    )

    doc.add_paragraph(
        '"CHINH SACH THUONG QIP CUA HWK" duoc cap nhat moi 6 thang. Phien ban dau tien duoc tao '
        'vao thang 2/2025, va tai lieu nay la Phien ban 3 da cap nhat, phan anh he thong V10 hien tai.'
    )

    doc.add_heading("MUC DICH", level=2)
    doc.add_paragraph(
        "Muc dich cua chuong trinh nay rat don gian:"
    )
    bullet(doc, "Tang chat luong san pham: Khi moi nguoi co dong luc, chat luong tot hon.")
    bullet(doc, "Thuong cho nguoi lam tot: Ai lam viec cham chi va chinh xac se duoc thuong xung dang.")
    bullet(doc, "Dat muc tieu 5Q: Giup HWK tro thanh nha may chat luong hang dau toan cau.")
    bullet(doc, "Giu chan nhan vien gioi: Nhan vien gioi se muon o lai khi duoc doi xu tot.")
    bullet(doc, "Tang tinh than dong doi: Khi ca nhom lam tot, ca nhom duoc thuong.")

    doc.add_paragraph()
    bold_para(doc, "Dieu quan trong nhat ban can nho: ",
              "Ban phai dat TAT CA cac dieu kien ap dung cho chuc vu cua ban. "
              "Chi can 1 dieu kien khong dat = khong nhan thuong thang do. "
              "Nhung dung lo — tai lieu nay se giai thich tung dieu kien rat chi tiet de ban hieu ro.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 2
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 2. GIAI THICH CHUNG VE CAU TRUC CHINH SACH THUONG", level=1)

    doc.add_paragraph(
        "Chinh sach thuong QIP cua HWK co 2 Chuong trinh chinh. "
        "Hay nghi nhu 2 cach de nhan thuong:"
    )
    bullet(doc, 'Chuong trinh 1: "Thuong hang thang" — Moi thang ban dat du dieu kien, ban duoc thuong.')
    bullet(doc, 'Chuong trinh 2: "Thuong Talent Pool 6 thang" — Thuong them cho nhan vien xuat sac duoc chon.')

    doc.add_paragraph()
    doc.add_paragraph(
        "Trong Chuong trinh 1, moi nhan vien QIP duoc xep vao 1 trong 3 loai (TYPE). "
        "Tuy theo loai, cach tinh thuong se khac nhau:"
    )

    doc.add_heading("Type-1: Nhan vien QIP truc tiep tai Assembly", level=2)
    doc.add_paragraph(
        "Day la nhung nguoi truc tiep kiem tra chat luong san pham tai khu vuc Assembly. "
        "Type-1 duoc danh gia khat khe nhat voi toi da 10 dieu kien chia thanh 3 nhom:"
    )
    bullet(doc, "Nhom Chuyen can (C1-C4): Di lam day du, khong vang khong phep, lam du ngay")
    bullet(doc, "Nhom AQL (C5-C8): Khong co loi chat luong AQL")
    bullet(doc, "Nhom 5PRS (C9-C10): Ket qua kiem tra 5PRS tot")

    doc.add_paragraph()
    add_example_box(doc, "Vi du de hieu:", [
        "Anh Nguyen la Assembly Inspector. Anh phai dat 8 dieu kien (C1-C6, C9, C10).",
        "Thang nay anh dat 7/8 dieu kien, chi truot C9 (ty le 5PRS thap).",
        "Ket qua: Anh Nguyen KHONG nhan thuong thang nay vi khong dat 100%.",
        "",
        "Chi Linh cung la Assembly Inspector. Chi dat 8/8 dieu kien.",
        "Ket qua: Chi Linh DUOC nhan thuong thang nay!"
    ])

    bold_para(doc, "QUAN TRONG: ",
              "Khong phai tat ca 10 dieu kien deu ap dung cho moi chuc vu. "
              "Vi du Assembly Inspector can 8 dieu kien, Line Leader chi can 5 dieu kien. "
              "Nhung BAT KE bao nhieu dieu kien, ban phai dat 100% dieu kien ap dung cho ban.")

    doc.add_heading("Type-2: Nhan vien QIP gian tiep", level=2)
    doc.add_paragraph(
        "Day la nhung nguoi khong truc tiep kiem tra tai Assembly nhung van dong gop vao chat luong. "
        "Type-2 chi can dat 4 dieu kien chuyen can (C1-C4)."
    )
    bullet(doc, "Ty le di lam >= 88%")
    bullet(doc, "Vang khong phep <= 2 ngay")
    bullet(doc, "Ngay lam thuc te > 0")
    bullet(doc, "Ngay lam toi thieu >= 12 ngay")

    doc.add_paragraph(
        "So tien thuong cua Type-2 duoc tinh dua tren muc thuong trung binh cua Type-1 tuong ung. "
        "Cac chuc vu quan ly (Group Leader, Supervisor, Manager) con co he so nhan them."
    )

    doc.add_heading("Type-3: Nhan vien QIP moi", level=2)
    doc.add_paragraph(
        "Neu ban moi vao lam va co duoi 30 ngay lam viec, ban thuoc Type-3. "
        "Type-3 chua duoc danh gia va chua nhan thuong. "
        "Sau khoang 3 thang, ban se duoc nang len Type-2 va bat dau nhan thuong."
    )

    add_example_box(doc, "Vi du:", [
        "Chi Hoa moi vao lam ngay 15/01/2026. Den het thang 2, chi moi lam duoc 35 ngay.",
        "Tu thang 3/2026, chi duoc nang len Type-2 va bat dau duoc danh gia thuong."
    ])

    doc.add_heading("Chuong trinh 2 — Thuong Talent Pool 6 thang", level=2)
    doc.add_paragraph(
        "Day la chuong trinh thuong bo sung cho nhan vien xuat sac duoc quan ly chon vao nhom Talent Pool. "
        "Thuong nay CONG THEM vao thuong Chuong trinh 1, khong thay the. "
        "Xem Phan 6 de biet chi tiet."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 3
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 3. PHAM VI VA DOI TUONG AP DUNG", level=1)

    doc.add_paragraph(
        "Chinh sach nay ap dung cho TAT CA thanh vien doi QIP. "
        "Ngoai tru nhan vien QIP Type-3 (nhan vien moi chua du 30 ngay lam viec)."
    )

    doc.add_heading("Phan loai nhan vien QIP", level=2)

    bold_para(doc, "Type 1 — QIP truc tiep tai Assembly: ", "kiem tra chat luong san pham cuoi cung")
    bullet(doc, "Kiem tra vien tai day chuyen Assembly (Assembly Inspector)")
    bullet(doc, "Kiem tra vien tai day chuyen Repacking")
    bullet(doc, "Kiem tra vien AQL (AQL Inspector)")
    bullet(doc, "Dao tao vien kiem tra (Doi Audit & Training)")
    bullet(doc, "Doi Model Master QIP")
    bullet(doc, "Kiem tra vien RQC Assembly (ma chuc vu A1B — mien C7 va C10 tu thang 2/2026)")
    bullet(doc, "Line Leader")

    bold_para(doc, "Type 2 — QIP gian tiep tai Assembly: ", "dong gop vao chat luong san pham cuoi cung")
    bullet(doc, "Doi QIP Vat lieu (MTL)")
    bullet(doc, "Doi QIP Cat (Cutting)")
    bullet(doc, "Doi QIP OSC/Inhouse")
    bullet(doc, "Doi QIP May (Stitching)")
    bullet(doc, "Doi QIP De (Bottom)")
    bullet(doc, "Doi QIP Dong goi (Packing)")
    bullet(doc, "Doi QA")
    bullet(doc, "Doi QIP van phong")
    bullet(doc, "Doi OCPT")
    bullet(doc, "Tat ca thanh vien QIP khac khong thuoc Type 1 va Bo phan Kiem nghiem (LAB)")

    bold_para(doc, "Type 3 — QIP moi: ", "Nhan vien moi voi duoi 30 ngay lam viec")
    bullet(doc, "Khong danh gia hoac thanh toan thuong trong thoi gian thu viec")
    bullet(doc, "Tu dong nang len Type-2 sau khoang 3 thang")

    doc.add_heading("Ky danh gia & Lich thanh toan", level=2)
    bullet(doc, "Ky danh gia: Tu ngay dau tien den ngay cuoi cung cua thang.")
    bullet(doc, "Ky xet thuong: Tu ngay 1 den ngay 15 cua thang sau.")
    bullet(doc, "Nhan vien dang lam viec: Thanh toan trong ky luong tiep theo.")
    bullet(doc, "Nhan vien da nghi: Thanh toan trong vong 14 ngay lam viec ke tu ngay nghi.")

    add_example_box(doc, "Vi du ve lich thanh toan:", [
        "Ban lam viec tot trong thang 3/2026.",
        "Tu ngay 1-15/4/2026, he thong xet thuong.",
        "Ban se nhan tien thuong trong ky luong thang 4/2026."
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 4
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 4. DINH NGHIA & THUAT NGU VIET TAT", level=1)

    doc.add_paragraph(
        "Duoi day la nhung tu viet tat va thuat ngu ban se gap trong tai lieu nay. "
        "Hay doc qua de khi gap lai ban se hieu ngay:"
    )

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Thuat ngu", "Nghia la gi?"])

    terms = [
        ["QIP", "Chuong trinh Cai tien Chat luong — He thong thuong cho nhan vien kiem tra chat luong cua HWK"],
        ["5Q", "5 Muc tieu Chat luong — Muc tieu chat luong ma HWK dang phan dau dat duoc"],
        ["AQL", "Muc Chat luong Chap nhan duoc — Tieu chuan kiem tra lay mau quoc te. "
                "Neu san pham bi loi AQL, nghia la san pham khong dat chuan chat luong"],
        ["5PRS", "Lay mau ngau nhien 5 doi — Cach kiem tra: lay ngau nhien 5 doi giay de kiem tra chat luong"],
        ["CFA", "Kiem tra vien Cuoi cung duoc Chung nhan — Chung chi AQL do adidas cap. "
                "Ai co chung chi nay se duoc thuong them"],
        ["PO", "Don hang San xuat — Moi don hang san xuat rieng le"],
        ["C1-C10", "10 dieu kien thuong. C1 la dieu kien 1, C2 la dieu kien 2, v.v."],
        ["TYPE-1", "Nhan vien truc tiep kiem tra chat luong (danh gia day du dieu kien)"],
        ["TYPE-2", "Nhan vien gian tiep/quan ly (chi danh gia chuyen can C1-C4)"],
        ["TYPE-3", "Nhan vien moi (chua du 30 ngay lam viec, chua nhan thuong)"],
        ["Bang luy tien", "Bang so tien thuong tang dan: Thang 1 duoc 150K, thang 2 duoc 250K, ... len toi 1 trieu"],
        ["Thang lien tuc", "So thang lien tiep ban dat du dieu kien. Neu that bai 1 thang, dem lai tu dau"],
        ["AR1", "Vang khong phep — loai vang mat bi tinh la 'khong phep'"],
        ["AR2", "Vang co phep vi om — loai vang mat KHONG bi tinh la 'khong phep'"],
        ["Allowance", "Ngoai le — Khi quan ly cho phep mien dieu kien trong truong hop dac biet"],
        ["Firestore", "Co so du lieu tren may cua Google, noi luu tat ca du lieu cua he thong"],
        ["RBAC", "He thong phan quyen — admin duoc lam nhieu hon nguoi dung thuong"],
    ]
    for t in terms:
        add_row(tbl, t)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 5. CHUONG TRINH 1
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 5. PHUONG PHAP TINH THUONG — CHUONG TRINH 1", level=1)

    doc.add_paragraph(
        "Phan nay giai thich cach tinh thuong hang thang. "
        "Day la phan quan trong nhat — hay doc ky de hieu cach he thong tinh thuong cho ban."
    )

    # 5.1
    doc.add_heading("5.1 Tong quan 10 Dieu kien Thuong", level=2)
    doc.add_paragraph(
        "He thong V10 danh gia toi da 10 dieu kien cho moi thanh vien QIP. "
        "Khong phai tat ca 10 dieu kien deu ap dung cho moi chuc vu — "
        "dieu kien ap dung phu thuoc vao chuc vu cua ban (xem Phan 5.2 va Phu luc 2). "
        "Nhung TAT CA dieu kien ap dung phai duoc dap ung (ty le dat 100%) de nhan thuong."
    )

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["#", "Dieu kien", "Nhom", "Nguong mac dinh", "Giai thich don gian"])

    conditions = [
        ["C1", "Ty le di lam", "Chuyen can", ">= 88%", "Ban phai di lam it nhat 88% so ngay lam viec. "
                                                          "Nghi phep hop le khong tinh la vang"],
        ["C2", "Vang khong phep", "Chuyen can", "<= 2 ngay", "Khong duoc vang khong phep qua 2 ngay/thang"],
        ["C3", "Ngay lam thuc te", "Chuyen can", "> 0 ngay", "Phai co it nhat 1 ngay di lam trong thang"],
        ["C4", "Ngay lam toi thieu", "Chuyen can", ">= 12 ngay", "Phai lam it nhat 12 ngay trong thang"],
        ["C5", "Loi AQL ca nhan", "AQL", "= 0 loi", "KHONG duoc co bat ky loi AQL nao. 1 loi = truot"],
        ["C6", "AQL lien tuc ca nhan", "AQL", "Khong lien tuc", "Khong duoc co loi AQL 2 thang lien tiep (2026+)"],
        ["C7", "AQL lien tuc nhom/KV", "AQL", "Khong lien tuc", "Nhan vien duoi quyen ban khong duoc loi lien tuc"],
        ["C8", "Ty le tu choi KV", "AQL", "< 3.0%", "Ty le tu choi cua khu vuc phu trach phai duoi 3%"],
        ["C9", "Ty le dat 5PRS", "5PRS", ">= 95%", "Kiem tra 5PRS phai dat tu 95% tro len"],
        ["C10", "SL kiem tra 5PRS", "5PRS", ">= 100 doi", "Phai kiem tra it nhat 100 doi giay"],
    ]
    for c in conditions:
        add_row(tbl, c)

    doc.add_paragraph()
    bold_para(doc, "Luu y ve nguong co the thay doi: ",
              "Cac nguong tren la gia tri mac dinh. Quan tri vien co the dieu chinh nguong hang thang "
              "qua Admin Panel. Moi thay doi deu duoc ghi lai day du.")

    doc.add_paragraph()

    # ---- DETAILED CALCULATION EXAMPLES ----
    doc.add_heading("5.1.1 Vi du tinh toan chi tiet — Dieu kien C1 (Ty le di lam)", level=3)

    doc.add_paragraph(
        "Day la dieu kien quan trong nhat va cung de bi nham nhat. "
        "Hay doc ky vi du sau:"
    )

    bold_para(doc, "Cong thuc tinh ty le di lam:")
    bullet(doc, "Buoc 1: Tinh Ngay vang = Tong ngay lam - Ngay lam thuc te - Ngay nghi phep duoc duyet")
    bullet(doc, "Buoc 2: Tinh Ty le vang = (Ngay vang / Tong ngay lam) x 100")
    bullet(doc, "Buoc 3: Tinh Ty le di lam = 100 - Ty le vang")

    doc.add_paragraph()

    add_example_box(doc, "Vi du 1: Chi Lan — DAT dieu kien C1", [
        "Thong tin thang 3/2026:",
        "- Tong ngay lam viec trong thang: 27 ngay",
        "- Ngay di lam thuc te: 23 ngay",
        "- Nghi phep nam: 2 ngay (duoc duyet)",
        "- Vang khong phep (AR1): 2 ngay",
        "",
        "Cach tinh:",
        "  Ngay vang = 27 - 23 - 2 = 2 ngay",
        "  Ty le vang = (2 / 27) x 100 = 7.4%",
        "  Ty le di lam = 100 - 7.4 = 92.6%",
        "",
        "Ket qua: 92.6% >= 88% => DAT dieu kien C1!"
    ])

    add_example_box(doc, "Vi du 2: Anh Hung — KHONG DAT dieu kien C1", [
        "Thong tin thang 3/2026:",
        "- Tong ngay lam viec trong thang: 27 ngay",
        "- Ngay di lam thuc te: 20 ngay",
        "- Nghi om AR2: 1 ngay (duoc duyet)",
        "- Vang khong phep (AR1): 6 ngay",
        "",
        "Cach tinh:",
        "  Ngay vang = 27 - 20 - 1 = 6 ngay",
        "  Ty le vang = (6 / 27) x 100 = 22.2%",
        "  Ty le di lam = 100 - 22.2 = 77.8%",
        "",
        "Ket qua: 77.8% < 88% => KHONG DAT dieu kien C1!",
        "Anh Hung se khong nhan thuong thang nay."
    ])

    add_example_box(doc, "Vi du 3: Chi Mai — nghi thai san", [
        "Thong tin thang 3/2026:",
        "- Tong ngay lam viec trong thang: 27 ngay",
        "- Ngay di lam thuc te: 10 ngay",
        "- Nghi thai san: 15 ngay (duoc duyet)",
        "- Vang khong phep (AR1): 2 ngay",
        "",
        "Cach tinh:",
        "  Ngay vang = 27 - 10 - 15 = 2 ngay",
        "  Ty le vang = (2 / 27) x 100 = 7.4%",
        "  Ty le di lam = 100 - 7.4 = 92.6%",
        "",
        "Ket qua: 92.6% >= 88% => DAT dieu kien C1!",
        "Nghi thai san duoc duyet nen KHONG tinh la vang."
    ])

    doc.add_paragraph()
    bold_para(doc, "Diem then chot: ",
              "Nghi phep duoc duyet (phep nam, thai san, om AR2, cong tac, v.v.) "
              "KHONG tinh la vang mat. Chi co AR1 (vang khong phep) moi tinh la vang.")

    doc.add_paragraph()

    # Detailed absence classification
    doc.add_heading("5.1.2 Phan loai vang mat — Cai nao tinh, cai nao khong?", level=3)

    doc.add_paragraph(
        "Day la danh sach day du cac loai vang mat. Ban can biet loai nao KHONG bi tinh "
        "va loai nao SE BI tinh vao ty le vang:"
    )

    bold_para(doc, "KHONG tinh vao ty le vang (Nghi phep duoc duyet):")
    bullet(doc, "Sinh thuong 1 con (Nghi thai san)")
    bullet(doc, "Phep nam")
    bullet(doc, "Vang co phep")
    bullet(doc, "Duong suc sinh thuong")
    bullet(doc, "Kham thai binh thuong")
    bullet(doc, "Con duoi 3 tuoi bi benh")
    bullet(doc, "AR2 - om ngan ngay")
    bullet(doc, "Di cong tac")
    bullet(doc, "Nghia vu quan su")
    bullet(doc, "Di lam khong quet the")
    bullet(doc, "Cong nhan vien moi")
    bullet(doc, "Nghi bu")

    bold_para(doc, "CO TINH vao ty le vang (Vang khong phep):")
    bullet(doc, "AR1 - Vang khong phep")
    bullet(doc, "AR1 - Gui thu")

    add_example_box(doc, "Nho don gian:", [
        "Chi co AR1 moi la 'vang khong phep'.",
        "TAT CA cac loai nghi khac (phep nam, thai san, om AR2, cong tac, v.v.) deu KHONG tinh.",
        "Nen neu ban nghi phep dung quy dinh, ty le di lam cua ban se khong bi anh huong!"
    ])

    doc.add_page_break()

    # 5.2
    doc.add_heading("5.2 Ma tran Ap dung Dieu kien theo Chuc vu (Type-1)", level=2)
    doc.add_paragraph(
        "Khong phai tat ca chuc vu deu phai dat du 10 dieu kien. "
        "Bang duoi day cho biet chuc vu cua ban can dat nhung dieu kien nao:"
    )

    tbl = doc.add_table(rows=1, cols=11)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Chuc vu", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9", "C10"])

    matrix = [
        ["ASSEMBLY INSPECTOR", "Co", "Co", "Co", "Co", "Co", "Co", "—", "—", "Co", "Co"],
        ["RQC ASSEMBLY INSPECTOR\n(A1B, tu 02/2026)", "Co", "Co", "Co", "Co", "Co", "Co", "—", "—", "Co", "—"],
        ["AQL INSPECTOR", "Co", "Co", "Co", "Co", "Co", "—", "—", "—", "—", "—"],
        ["LINE LEADER", "Co", "Co", "Co", "Co", "—", "—", "Co", "—", "—", "—"],
        ["AUDITOR & TRAINER", "Co", "Co", "Co", "Co", "—", "—", "Co", "Co", "—", "—"],
        ["MODEL MASTER", "Co", "Co", "Co", "Co", "—", "—", "—", "Co", "—", "—"],
        ["Quan ly (Group Leader,\nSupervisor, Manager, v.v.)", "Co", "Co", "Co", "Co", "—", "—", "—", "—", "—", "—"],
    ]
    for m in matrix:
        add_row(tbl, m)

    doc.add_paragraph()
    doc.add_paragraph(
        'Ghi chu: "Co" = dieu kien ap dung cho chuc vu nay. "—" = khong ap dung (duoc mien).'
    )

    add_example_box(doc, "Cach doc bang nay:", [
        "Neu ban la ASSEMBLY INSPECTOR, ban can dat: C1, C2, C3, C4, C5, C6, C9, C10 = 8 dieu kien.",
        "Neu ban la LINE LEADER, ban can dat: C1, C2, C3, C4, C7 = 5 dieu kien.",
        "Neu ban la AQL INSPECTOR, ban can dat: C1, C2, C3, C4, C5 = 5 dieu kien.",
        "Quan ly (GL, Supervisor, Manager) chi can dat: C1, C2, C3, C4 = 4 dieu kien."
    ])

    bold_para(doc, "RQC Assembly Inspector (A1B): ",
              "Tu thang 2/2026, ma chuc vu A1B duoc phan loai la RQC_ASSEMBLY_INSPECTOR. "
              "C7 (AQL nhom/KV) va C10 (SL kiem tra 5PRS) duoc mien vi kiem tra vien RQC "
              "thuc hien kiem tra quy trinh va bao cao, khong phai kiem tra day chuyen.")

    doc.add_page_break()

    # 5.3
    doc.add_heading("5.3 Bang Thuong Luy tien Type-1", level=2)
    doc.add_paragraph(
        "Day la phan thu vi nhat! So tien thuong TANG DAN theo so thang lien tuc ban dat du dieu kien. "
        "Cang lam tot lau, tien thuong cang cao:"
    )

    tbl = doc.add_table(rows=2, cols=14)
    tbl.style = 'Table Grid'
    headers = ["Thang", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12-15"]
    amounts = ["VND (x1.000)", "150", "250", "300", "350", "400", "450", "500", "650", "750", "850", "950", "1.000"]

    for i, text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")
    for i, text in enumerate(amounts):
        cell = tbl.rows[1].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8)
        if i == 0:
            run.bold = True

    doc.add_paragraph()

    add_example_box(doc, "Vi du thuc te — Hanh trinh thuong cua Anh Tuan:", [
        "Thang 1/2026: Dat du dieu kien -> Thang lien tuc = 1 -> Thuong 150,000 VND",
        "Thang 2/2026: Dat du dieu kien -> Thang lien tuc = 2 -> Thuong 250,000 VND",
        "Thang 3/2026: Dat du dieu kien -> Thang lien tuc = 3 -> Thuong 300,000 VND",
        "Thang 4/2026: KHONG DAT (vang 3 ngay khong phep) -> Reset ve 0!",
        "Thang 5/2026: Dat du dieu kien -> Thang lien tuc = 1 -> Thuong 150,000 VND (bat dau lai tu dau)",
        "Thang 6/2026: Dat du dieu kien -> Thang lien tuc = 2 -> Thuong 250,000 VND",
        "",
        "Neu anh Tuan khong bi truot thang 4, den thang 6 anh da nhan 400,000 VND thay vi 250,000 VND!",
        "=> Bai hoc: Giu lien tuc rat quan trong vi 1 thang truot se mat nhieu tien thuong."
    ])

    doc.add_paragraph("Quy tac chinh:")
    bullet(doc, "Gioi han toi da: 1,000,000 VND/thang (dat tai thang thu 12 lien tuc)")
    bullet(doc, "Reset khi that bai: Neu ban khong dat dieu kien bat ky thang nao, thang lien tuc RESET ve 0")
    bullet(doc, "Chuyen tiep: Thang lien tuc duoc tinh tu thang truoc khi dat dieu kien")
    bullet(doc, "Theo doi toi da: 15 thang (so tien giu nguyen 1,000,000 VND tu thang 12 tro di)")

    doc.add_page_break()

    # 5.4
    doc.add_heading("5.4 AQL Inspector — Tinh thuong 3 Phan", level=2)
    doc.add_paragraph(
        "AQL Inspector duoc tinh thuong dac biet gom 3 phan. "
        "Neu ban la AQL Inspector, so tien thuong hang thang la TONG cua ca 3 phan:"
    )

    bold_para(doc, "Phan 1: Ket qua Danh gia Kiem tra AQL")
    doc.add_paragraph(
        "Dua tren bang luy tien (giong Phan 5.3). Ngoai ra, ty le tu choi boi adidas/T1QM/ben thu 3 phai < 3%."
    )

    bold_para(doc, "Phan 2: Thuong Chung chi CFA")
    doc.add_paragraph(
        "Neu ban co chung chi AQL do adidas cap (CFA), ban duoc them co dinh 700,000 VND/thang. "
        "Khong co chung chi = khong co Phan 2."
    )

    bold_para(doc, "Phan 3: Thuong Phong ngua Khieu nai HWK")
    doc.add_paragraph("Thuong bo sung cho viec phong ngua khieu nai, tang luy tien tu thang thu 4:")

    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = 'Table Grid'
    headers2 = ["Thang", "1-3", "4-6", "7-9", "10-12", "13-15"]
    amounts2 = ["VND (x1.000)", "0", "300", "500", "700", "900"]
    for i, text in enumerate(headers2):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")
    for i, text in enumerate(amounts2):
        cell = tbl.rows[1].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True

    doc.add_paragraph()

    add_example_box(doc, "Vi du: AQL Inspector co CFA tai thang lien tuc thu 12:", [
        "Phan 1 (Luy tien):           1,000,000 VND",
        "Phan 2 (Chung chi CFA):        700,000 VND",
        "Phan 3 (Phong ngua HWK):       700,000 VND (thang 10-12)",
        "---------------------------------------",
        "TONG CONG:                    2,400,000 VND/thang!",
        "",
        "Day la muc thuong cao nhat ma AQL Inspector co the nhan duoc."
    ])

    add_example_box(doc, "Vi du: AQL Inspector KHONG co CFA tai thang lien tuc thu 5:", [
        "Phan 1 (Luy tien):             400,000 VND (thang 5)",
        "Phan 2 (Chung chi CFA):              0 VND (khong co CFA)",
        "Phan 3 (Phong ngua HWK):       300,000 VND (thang 4-6)",
        "---------------------------------------",
        "TONG CONG:                      700,000 VND/thang"
    ])

    doc.add_page_break()

    # 5.5
    doc.add_heading("5.5 Tinh Thuong Line Leader (Type-1)", level=2)
    doc.add_paragraph(
        "Line Leader khong tinh theo bang luy tien. Thay vao do, thuong cua Line Leader "
        "phu thuoc vao ket qua cua nhung nhan vien duoi quyen:"
    )
    bold_para(doc, "Cong thuc: ", "Tong thuong nhan vien duoi quyen x 12%")

    add_example_box(doc, "Vi du:", [
        "Line Leader Anh Duc co 10 nhan vien duoi quyen.",
        "8 nguoi dat dieu kien va duoc thuong, tong cong = 2,500,000 VND",
        "2 nguoi khong dat va khong duoc thuong = 0 VND",
        "",
        "Thuong cua Anh Duc = 2,500,000 x 12% = 300,000 VND",
        "",
        "Luu y: Anh Duc cung phai dat tat ca dieu kien ap dung (C1, C2, C3, C4, C7).",
        "Neu anh Duc khong dat C1, anh se khong nhan thuong du nhan vien duoi quyen dat tot."
    ])

    # 5.6
    doc.add_heading("5.6 Tinh Thuong Type-2", level=2)
    doc.add_paragraph(
        "Thanh vien Type-2 phai dat 100% dieu kien chuyen can (C1-C4). "
        "So tien thuong duoc tinh dua tren muc thuong trung binh cua cac chuc vu Type-1 tuong ung."
    )

    bold_para(doc, "Bang Anh xa Chuc vu TYPE-2:")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Chuc vu Type-2", "Tham chieu thuong", "Giai thich"])

    mappings = [
        ["BOTTOM INSPECTOR", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["CUTTING INSPECTOR", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["MTL INSPECTOR", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["STITCHING INSPECTOR", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["OSC INSPECTOR", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["OCPT STAFF", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["RQC", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["AQL INSPECTOR (TYPE-2)", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["QA TEAM (mac dinh)", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["QA TEAM (ma QA3A)", "TB GROUP LEADER", "Dac biet: anh xa toi TB Group Leader"],
        ["QA TEAM (ma QA3B)", "TB ASSEMBLY INSPECTOR", "Lay thuong trung binh cua Assembly Inspector Type-1"],
        ["LINE LEADER (TYPE-2)", "TB LINE LEADER", "Lay thuong trung binh cua Line Leader Type-1"],
    ]
    for m in mappings:
        add_row(tbl, m)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("TB = Trung binh")
    run.italic = True

    add_example_box(doc, "Vi du:", [
        "Chi Thanh la MTL Inspector (Type-2). Chi dat du 4 dieu kien chuyen can.",
        "Thang nay, trung binh thuong cua Assembly Inspector Type-1 la 350,000 VND.",
        "=> Chi Thanh se nhan 350,000 VND."
    ])

    # 5.7
    doc.add_heading("5.7 He so Nhan Chuc vu Quan ly Type-2", level=2)
    doc.add_paragraph(
        "Cac chuc vu quan ly trong Type-2 duoc tinh thuong = "
        "TB thuong LINE LEADER Type-1 (chi tinh nguoi nhan thuong) x he so nhan:"
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Chuc vu", "He so", "Cong thuc", "Giai thich"])

    multipliers = [
        ["GROUP LEADER", "2.0x", "TB LL x 2.0", "TB Line Leader nhan thuong x 2.0"],
        ["SUPERVISOR /\n(V) SUPERVISOR /\nA.SUPERVISOR", "2.5x", "TB LL x 2.5", "TB Line Leader nhan thuong x 2.5"],
        ["A.MANAGER", "3.0x", "TB LL x 3.0", "TB Line Leader nhan thuong x 3.0"],
        ["MANAGER", "3.5x", "TB LL x 3.5", "TB Line Leader nhan thuong x 3.5"],
        ["S.MANAGER", "4.0x", "TB LL x 4.0", "TB Line Leader nhan thuong x 4.0"],
    ]
    for m in multipliers:
        add_row(tbl, m)

    doc.add_paragraph()

    add_example_box(doc, "Vi du: Neu TB thuong Line Leader Type-1 la 400,000 VND:", [
        "Group Leader   = 400,000 x 2.0 = 800,000 VND",
        "(V) Supervisor = 400,000 x 2.5 = 1,000,000 VND",
        "A.Manager      = 400,000 x 3.0 = 1,200,000 VND",
        "Manager        = 400,000 x 3.5 = 1,400,000 VND",
        "S.Manager      = 400,000 x 4.0 = 1,600,000 VND"
    ])

    bold_para(doc, "Luu y: ",
              '"TB LINE LEADER nhan thuong" nghia la trung binh thuong cua Line Leader Type-1 '
              "chi tinh nhung nguoi NHAN thuong (> 0 VND). Nguoi co 0 VND bi loai khoi phep tinh trung binh.")

    # 5.8
    doc.add_heading("5.8 Thuong Doi Audit & Training (Trainer)", level=2)
    doc.add_paragraph(
        "Moi nhan vien dao tao co khu vuc phu trach rieng de kiem tra va dao tao kiem tra vien."
    )
    bullet(doc, "Dieu kien ap dung: C1, C2, C3, C4, C7 (AQL lien tuc nhom/KV), C8 (Ty le tu choi KV)")
    bullet(doc, "Neu ty le loi (Ty le Tu choi Khu vuc, C8) duoi nguong 3%: Dat yeu cau.")
    bullet(doc, "Neu ty le loi tren nguong: Khong dat yeu cau.")
    bullet(doc, "Bang luy tien (Phan 5.3) ap dung de tinh so tien thuong.")

    add_example_box(doc, "Vi du:", [
        "Chi Hien la Trainer phu trach toa nha A2.",
        "Thang nay ty le tu choi toa nha A2 la 2.1% (< 3%) => C8 DAT.",
        "Kiem tra nhan vien duoi quyen: khong co loi AQL lien tuc => C7 DAT.",
        "Chuyen can tot: C1, C2, C3, C4 deu DAT.",
        "Ket qua: Chi Hien nhan thuong thang nay theo bang luy tien."
    ])

    # 5.9
    doc.add_heading("5.9 Thuong Doi Model Master", level=2)
    doc.add_paragraph("Thuong doi Model Master dua tren ty le tu choi AQL HWK tren TAT CA nha may.")
    bullet(doc, "Neu ty le tu choi AQL chinh thuc cua HWK tren tat ca nha may (C8) duoi nguong: Du dieu kien.")
    bullet(doc, "Dieu kien ap dung: C1, C2, C3, C4, C8 (Ty le Tu choi Khu vuc)")
    bullet(doc, "Bang luy tien (Phan 5.3) ap dung de tinh so tien thuong.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 6
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 6. PHUONG PHAP TINH THUONG — CHUONG TRINH 2", level=1)

    doc.add_paragraph(
        'Chuong trinh 2 la "Chuong trinh Thuong Talent Pool 6 Thang." '
        "Chi nhung nhan vien xuat sac duoc quan ly chon moi duoc tham gia."
    )

    bold_para(doc, "QUAN TRONG: ",
              "Ai nhan thuong theo Chuong trinh 1 cung co the nhan thuong theo Chuong trinh 2. "
              "Hai chuong trinh CONG DON voi nhau, khong thay the.")

    doc.add_paragraph()
    doc.add_paragraph("Quy trinh Chuong trinh 2 gom 6 buoc:")

    steps = [
        ("Buoc 1: Thiet lap Muc tieu",
         "Moi thang 1 va thang 7, doi QIP danh gia hieu suat 6 thang truoc va lap ke hoach 6 thang toi."),
        ("Buoc 2: Phan cong Nhiem vu",
         "Quan ly QIP chon nguoi phu hop cho moi muc tieu. So tien thuong va ky thanh toan duoc xac dinh. COO phe duyet."),
        ("Buoc 3: Theo doi Hang thang",
         "Doi QA, doi van phong va quan ly QIP danh gia hang thang bang KPI do luong duoc."),
        ("Buoc 4: Danh gia Cuoi cung",
         "Sau 6 thang, quan ly QIP xem xet ke hoach va ket qua, ra quyet dinh cuoi cung."),
        ("Buoc 5: Xac nhan Talent Pool",
         "Quan ly QIP yeu cau phe duyet tu HR va COO qua he thong groupware HS."),
        ("Buoc 6: Phat trien Talent Pool",
         "Quy trinh lien tuc moi 6 thang. Muc tieu chung nhan 3-5% tong thanh vien QIP. Chung nhan co hieu luc 12 thang."),
    ]
    for title, desc in steps:
        bold_para(doc, title)
        doc.add_paragraph(desc)

    doc.add_paragraph()
    bold_para(doc, "Trong he thong V10:")
    bullet(doc, "Tu dong ap dung: Thuong Talent Pool tu dong ap dung cung thuong thuong xuyen")
    bullet(doc, "Cong don: Thuong Talent Pool cong them vao thuong Chuong trinh 1")
    bullet(doc, "Thoi gian thanh toan: Cung ky thuong hang thang")
    bullet(doc, "Yeu cau dieu kien: Khong — thanh vien Talent Pool nhan thuong bat ke dat/khong dat dieu kien")
    bullet(doc, "Loai tru: Nhan vien da nghi tu dong bi loai")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 7
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 7. TOM TAT CAC DIEU KHOAN NGOAI LE", level=1)

    doc.add_heading("7.1 Cac truong hop KHONG nhan Thuong", level=2)

    doc.add_paragraph(
        "Hay doc ky phan nay de biet nhung truong hop nao ban se KHONG nhan thuong:"
    )

    bold_para(doc, "Doi voi Type-1 (Chuong trinh 1):")
    bullet(doc, "Neu bat ky dieu kien ap dung nao khong dat (yeu cau dat 100%), khong nhan thuong thang do.")
    bullet(doc, "Neu ty le di lam < 88%, hoac vang khong phep > 2 ngay, hoac ngay lam thuc te = 0, "
           "hoac ngay lam toi thieu < 12 ngay.")
    bullet(doc, "Neu co bat ky loi AQL nao (doi voi chuc vu ap dung C5).")
    bullet(doc, "Neu ket qua 5PRS thap hon nguong (doi voi chuc vu ap dung C9, C10).")
    bullet(doc, "Neu co khieu nai chinh thuc tu adidas va trach nhiem thuoc ve ca nhan.")
    bullet(doc, "Doi voi Line Leader: Neu nhan vien duoi quyen co loi AQL lien tuc vuot nguong (C7).")
    bullet(doc, "Doi voi Trainer: Neu ty le tu choi toa nha phu trach vuot 3% (C8).")

    bold_para(doc, "Doi voi Type-2 (Chuong trinh 1):")
    bullet(doc, "Neu dieu kien chuyen can (C1-C4) khong dat 100%, khong nhan thuong.")
    bullet(doc, "Neu co khieu nai chinh thuc tu adidas thuoc trach nhiem ca nhan.")

    bold_para(doc, "Doi voi Type-3 (Chuong trinh 1):")
    bullet(doc, "TAT CA thanh vien QIP Type-3 khong nhan thuong (nhan vien moi).")

    bold_para(doc, "Doi voi Chuong trinh 2:")
    bullet(doc, "Neu de xuat cua quan ly QIP bi COO tu choi, Chuong trinh 2 khong co hieu luc.")

    doc.add_heading("7.2 Truong hop Ngoai le DUOC nhan Thuong (Allowance)", level=2)

    doc.add_paragraph(
        "Trong mot so truong hop dac biet, quan ly QIP co the cho phep ngoai le (goi la 'Allowance'). "
        "Noi don gian: neu ban khong dat dieu kien vi ly do chinh dang, quan ly co the 'mien' cho ban."
    )

    doc.add_paragraph("Cac truong hop co the xin ngoai le:")
    bullet(doc, "Om dau, nhap vien, chan thuong bat ngo khien ban khong du ngay lam.")
    bullet(doc, "Loi AQL do loi cua doi san xuat, khong phai loi cua ban.")
    bullet(doc, "Ket qua 5PRS thap vi ban moi duoc chuyen khu vuc lam viec.")
    bullet(doc, "Cac truong hop dac biet khac duoc quan ly xem xet.")

    bold_para(doc, "4 ma ly do Allowance trong he thong:")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Ma ly do", "Nghia la gi?"])
    codes = [
        ["MEDICAL", "Ly do y te (om dau, chan thuong, nhap vien)"],
        ["COMPANY_ORDER", "Lenh cong ty (dieu chuyen, nhiem vu bat buoc)"],
        ["NATURAL_DISASTER", "Thien tai (lu lut, bao, v.v.)"],
        ["OTHER", "Ly do khac (yeu cau nhap chi tiet)"],
    ]
    for c in codes:
        add_row(tbl, c)

    doc.add_paragraph()
    doc.add_paragraph("Quy trinh Allowance:")
    bullet(doc, "1. Quan tri vien tim kiem nhan vien trong tab Allowance")
    bullet(doc, "2. He thong hien thi trang thai dieu kien hien tai (dat/khong dat)")
    bullet(doc, "3. Quan tri vien chon cac dieu kien khong dat can mien")
    bullet(doc, "4. Quan tri vien nhap ma ly do + chi tiet giai trinh (bat buoc)")
    bullet(doc, "5. Man hinh xem truoc hien thi tac dong truoc khi ap dung")
    bullet(doc, "6. Khi Ap dung: he thong cap nhat du lieu, tinh lai thuong, cap nhat tom tat")
    bullet(doc, "7. Allowance co the thu hoi (khoi phuc trang thai ban dau)")
    bullet(doc, "8. Tinh lai hang loat co san cho tat ca nhan vien allowance co thuong 0")

    doc.add_paragraph(
        "Tat ca ngoai le allowance deu duoc ghi lai day du: ngay, quan tri vien, "
        "nhan vien, dieu kien mien, ly do, chi tiet, va so tien thuong tinh lai."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHAN 8 — WITH SCREENSHOTS
    # ══════════════════════════════════════════
    doc.add_heading("PHAN 8. HE THONG & BANG DIEU KHIEN (V10) — HINH ANH MINH HOA", level=1)

    doc.add_paragraph(
        "He thong Thuong QIP HWK V10 la nen tang dam may tu dong hoa toan bo quy trinh tinh thuong. "
        "Duoi day la huong dan su dung voi hinh anh minh hoa."
    )

    doc.add_heading("8.1 Cach truy cap he thong", level=2)
    doc.add_paragraph(
        "De xem ket qua thuong cua ban, hay truy cap he thong theo cac buoc sau:"
    )
    bullet(doc, "Buoc 1: Mo trinh duyet web (Chrome, Safari, v.v.)")
    bullet(doc, "Buoc 2: Truy cap dia chi: https://moonkaicuzui.github.io/hwk-qip-incentive-v10/")
    bullet(doc, "Buoc 3: Dang nhap bang email va mat khau duoc cap")
    bullet(doc, "Buoc 4: Chon thang/nam can xem")
    bullet(doc, "Buoc 5: Xem ket qua tren bang dieu khien")

    doc.add_paragraph()
    doc.add_paragraph("Duoi day la man hinh chon thang/nam:")
    add_screenshot(doc, "02_selector_en.png",
                   "Hinh 1: Man hinh chon thang/nam — Chon thang va nam ban muon xem ket qua")

    doc.add_heading("8.2 Bang dieu khien tong quan", level=2)
    doc.add_paragraph(
        "Sau khi chon thang, ban se thay bang dieu khien tong quan voi cac thong tin chinh:"
    )
    bullet(doc, "So nguoi nhan thuong / tong so nguoi")
    bullet(doc, "Ty le thanh toan")
    bullet(doc, "Tong so tien thuong")
    bullet(doc, "Bieu do phan loai theo TYPE")
    bullet(doc, "Bieu do dieu kien — cho thay dieu kien nao nhieu nguoi truot nhat")

    doc.add_paragraph()
    add_screenshot(doc, "07_dashboard_summary_vi.png",
                   "Hinh 2: Bang dieu khien tong quan (tieng Viet) — Hien thi KPI va thong ke chinh")

    doc.add_heading("8.3 Xem chi tiet ca nhan", level=2)
    doc.add_paragraph(
        "Ban co the bam vao ten cua minh de xem chi tiet tung dieu kien:"
    )
    bullet(doc, "Trang thai tung dieu kien (Dat / Khong Dat)")
    bullet(doc, "Gia tri cu the cua tung dieu kien")
    bullet(doc, "So tien thuong duoc tinh")
    bullet(doc, "So thang lien tuc")

    doc.add_paragraph()
    add_screenshot(doc, "05_employee_modal_en.png",
                   "Hinh 3: Chi tiet ca nhan — Bam vao ten de xem ket qua tung dieu kien")

    doc.add_heading("8.4 Luong Du lieu & Pipeline CI/CD", level=2)
    doc.add_paragraph("He thong tu dong chay moi 6 gio (va kich hoat thu cong):")
    bullet(doc, "Google Drive: Tai len du lieu nguon (nhan su, chuyen can, AQL, 5PRS)")
    bullet(doc, "GitHub Actions: Pipeline CI/CD tu dong tai du lieu, dong bo nguong va cau hinh")
    bullet(doc, "Tu dong phat hien: He thong tu dong phat hien thang can tinh lai")
    bullet(doc, "Python Engine: Xu ly tat ca 10 dieu kien va tinh so tien thuong")
    bullet(doc, "Firestore Upload: Tai ket qua len co so du lieu Firestore")
    bullet(doc, "Xu ly Email: Gui thong bao email dang cho qua SMTP")
    bullet(doc, "Trien khai: Bang dieu khien duoc trien khai len GitHub Pages")

    doc.add_heading("8.5 Tinh nang Bang dieu khien (8 Tab)", level=2)
    bullet(doc, "Tab Tom tat: The KPI, bang phan loai TYPE, bieu do dieu kien, bieu do xu huong")
    bullet(doc, "Tab Chi tiet Chuc vu: Bang phan tich theo tung chuc vu")
    bullet(doc, "Tab Chi tiet Ca nhan: Danh sach nhan vien co tim kiem/loc/phan trang")
    bullet(doc, "Tab Tieu chi Thuong: Tham chieu nguong hien tai va mo ta dieu kien")
    bullet(doc, "Tab So do To chuc: Cau truc to chuc doi QIP")
    bullet(doc, "Tab Quan ly Nhom: Phan tich hieu suat cap nhom")
    bullet(doc, "Tab Xac minh: 12 the xac minh KPI, kiem tra tinh toan ven du lieu")
    bullet(doc, "Tab Tra cuu Chuyen can: Lich chuyen can ca nhan voi tra cuu hang ngay")

    doc.add_heading("8.6 Admin Panel (5 Tab)", level=2)
    doc.add_paragraph("Admin Panel chi danh cho quan tri vien:")
    bullet(doc, "Nguong: Dieu chinh 7 nguong hang thang voi lich su thay doi")
    bullet(doc, "Quan ly Cau hinh: Anh xa chuc vu, cau hinh bang luy tien")
    bullet(doc, "He thong: Giam sat trang thai pipeline, cau hinh he thong")
    bullet(doc, "Tra cuu Du lieu: Truy van va xac minh du lieu Firestore truc tiep")
    bullet(doc, "Allowances: Quan ly ngoai le voi quy trinh ap dung/thu hoi")

    doc.add_heading("8.7 Quan ly Nguong", level=2)
    doc.add_paragraph(
        "Tat ca gia tri nguong duoc quan ly tap trung trong Firestore. "
        "Moi thay doi tao ban ghi khong the sua doi bao gom: truong thay doi, gia tri cu/moi, "
        "nguoi thay doi (email), va thoi gian."
    )

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Thong so", "Mac dinh", "Cau hinh"])
    params = [
        ["Ty le di lam (C1)", "88%", "Co"],
        ["Vang khong phep (C2)", "2 ngay", "Co"],
        ["Ngay lam toi thieu (C4)", "12 ngay", "Co"],
        ["Ty le tu choi KV (C8)", "3.0%", "Co"],
        ["Ty le dat 5PRS (C9)", "95%", "Co"],
        ["SL toi thieu 5PRS (C10)", "100 doi", "Co"],
        ["Thang AQL lien tuc (C6, C7)", "3 thang", "Co"],
    ]
    for p_row in params:
        add_row(tbl, p_row)

    doc.add_heading("8.8 He thong Phan hoi", level=2)
    doc.add_paragraph(
        "Neu ban phat hien loi hoac co y kien dong gop, ban co the gui phan hoi qua he thong:"
    )
    bullet(doc, "Loai phan hoi: LOI (BUG), CAI TIEN, TINH NANG MOI, UI/UX, DU LIEU, KHAC")
    bullet(doc, "Muc uu tien: Khan cap, Cao, Trung binh, Thap")
    bullet(doc, "Luong trang thai: DA GUI -> DANG XEM XET -> DANG TIEN HANH -> HOAN THANH (hoac TU CHOI)")
    bullet(doc, "Tinh nang: Dinh kem hinh anh (toi da 3), nhieu nguoi nhan thong bao, tra loi admin qua email")

    doc.add_heading("8.9 He thong Thong bao Email", level=2)
    doc.add_paragraph("3 thong bao email tu dong:")
    bullet(doc, "Phan hoi moi -> Tat ca quan tri vien nhan email thong bao (3 ngon ngu: ko/en/vi)")
    bullet(doc, "Trang thai phan hoi thay doi -> Nguoi bao cao nhan thong bao cap nhat")
    bullet(doc, "Admin tra loi phan hoi -> Nguoi bao cao nhan email tra loi")

    doc.add_heading("8.10 Ho tro Da ngon ngu", level=2)
    doc.add_paragraph(
        "Toan bo bang dieu khien ho tro 3 ngon ngu: Tieng Han (ko), Tieng Anh (en), va Tieng Viet (vi). "
        "Ban co the chuyen doi ngon ngu bat cu luc nao."
    )

    doc.add_heading("8.11 Xac thuc & RBAC", level=2)
    doc.add_paragraph("Xac thuc Firebase Email/Mat khau:")
    bullet(doc, "Nguoi dung thuong: Xem bang dieu khien chi doc + gui phan hoi")
    bullet(doc, "Quan tri vien: Toan quyen bao gom admin panel, quan ly nguong, allowance, phan hoi")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHU LUC 1 — MASSIVELY EXPANDED
    # ══════════════════════════════════════════
    doc.add_heading("PHU LUC 1: 10 DIEU KIEN THUONG — GIAI THICH CHI TIET VOI VI DU", level=1)

    doc.add_paragraph(
        "Day la phan quan trong nhat cua tai lieu. Moi dieu kien se duoc giai thich ky voi vi du thuc te "
        "de ban hieu ro va biet cach dat thuong."
    )

    # ---- C1 ----
    doc.add_heading("DIEU KIEN 1 (C1): TY LE DI LAM", level=2)
    bold_para(doc, "Nguong: ", ">= 88% (mac dinh, co the thay doi boi admin)")
    bold_para(doc, "Ap dung cho: ", "TAT CA chuc vu Type-1 va Type-2")

    doc.add_paragraph(
        "Day la dieu kien co ban nhat: ban phai di lam du. "
        "He thong tinh ty le di lam cua ban theo cong thuc sau:"
    )

    bold_para(doc, "Cong thuc:")
    bullet(doc, "Ngay vang = Tong ngay lam - Ngay lam thuc te - Ngay nghi phep duoc duyet")
    bullet(doc, "Ty le vang = (Ngay vang / Tong ngay lam) x 100%")
    bullet(doc, "Ty le di lam = 100% - Ty le vang")

    doc.add_paragraph()
    bold_para(doc, "Diem quan trong nhat: ",
              "Nghi phep duoc duyet (phep nam, thai san, om AR2, cong tac, nghi bu, v.v.) "
              "KHONG tinh la vang mat. Chi co AR1 (vang khong phep) moi tinh la vang.")

    doc.add_paragraph()

    add_example_box(doc, "Vi du 1 — Chi Lan (DAT):", [
        "Tong ngay lam viec: 27 ngay",
        "Ngay di lam thuc te: 23 ngay",
        "Nghi phep nam (duoc duyet): 2 ngay",
        "Vang khong phep AR1: 2 ngay",
        "",
        "Tinh: Ngay vang = 27 - 23 - 2 = 2 ngay",
        "      Ty le vang = (2/27) x 100 = 7.4%",
        "      Ty le di lam = 100 - 7.4 = 92.6%",
        "",
        "92.6% >= 88% => DAT!"
    ])

    add_example_box(doc, "Vi du 2 — Anh Hung (KHONG DAT):", [
        "Tong ngay lam viec: 27 ngay",
        "Ngay di lam thuc te: 20 ngay",
        "Nghi om AR2 (duoc duyet): 1 ngay",
        "Vang khong phep AR1: 6 ngay",
        "",
        "Tinh: Ngay vang = 27 - 20 - 1 = 6 ngay",
        "      Ty le vang = (6/27) x 100 = 22.2%",
        "      Ty le di lam = 100 - 22.2 = 77.8%",
        "",
        "77.8% < 88% => KHONG DAT!",
        "Anh Hung se khong nhan thuong thang nay."
    ])

    add_example_box(doc, "Vi du 3 — Chi Mai nghi thai san (DAT):", [
        "Tong ngay lam viec: 27 ngay",
        "Ngay di lam thuc te: 10 ngay",
        "Nghi thai san (duoc duyet): 15 ngay",
        "Vang khong phep AR1: 2 ngay",
        "",
        "Tinh: Ngay vang = 27 - 10 - 15 = 2 ngay",
        "      Ty le vang = (2/27) x 100 = 7.4%",
        "      Ty le di lam = 100 - 7.4 = 92.6%",
        "",
        "92.6% >= 88% => DAT!",
        "Nghi thai san la nghi phep duoc duyet nen KHONG tinh la vang."
    ])

    add_example_box(doc, "Vi du 4 — Anh Minh di cong tac nhieu (DAT):", [
        "Tong ngay lam viec: 27 ngay",
        "Ngay di lam thuc te: 15 ngay",
        "Di cong tac (duoc duyet): 10 ngay",
        "Vang khong phep AR1: 2 ngay",
        "",
        "Tinh: Ngay vang = 27 - 15 - 10 = 2 ngay",
        "      Ty le vang = (2/27) x 100 = 7.4%",
        "      Ty le di lam = 100 - 7.4 = 92.6%",
        "",
        "92.6% >= 88% => DAT!",
        "Di cong tac cung la nghi phep duoc duyet."
    ])

    doc.add_paragraph()
    bold_para(doc, "Danh sach cac loai nghi KHONG tinh la vang:")
    bullet(doc, "Sinh thuong 1 con (Thai san)")
    bullet(doc, "Phep nam")
    bullet(doc, "Vang co phep")
    bullet(doc, "Duong suc sinh thuong")
    bullet(doc, "Kham thai binh thuong")
    bullet(doc, "Con duoi 3 tuoi bi benh")
    bullet(doc, "AR2 - om ngan ngay")
    bullet(doc, "Di cong tac")
    bullet(doc, "Nghia vu quan su")
    bullet(doc, "Di lam khong quet the")
    bullet(doc, "Cong nhan vien moi")
    bullet(doc, "Nghi bu")

    bold_para(doc, "Danh sach cac loai nghi CO TINH la vang:")
    bullet(doc, "AR1 - Vang khong phep")
    bullet(doc, "AR1 - Gui thu")

    doc.add_page_break()

    # ---- C2 ----
    doc.add_heading("DIEU KIEN 2 (C2): VANG KHONG PHEP", level=2)
    bold_para(doc, "Nguong: ", "<= 2 ngay (mac dinh)")
    bold_para(doc, "Ap dung cho: ", "TAT CA chuc vu Type-1 va Type-2")

    doc.add_paragraph(
        "Dieu kien nay dem so ngay ban vang khong phep (AR1) trong thang. "
        "Neu ban vang qua 2 ngay khong phep, ban se khong nhan thuong."
    )

    bold_para(doc, "Chi co 2 loai AR1:")
    bullet(doc, '"Vang khong phep" (AR1)')
    bullet(doc, '"AR1 - Gui thu" (gui thu xin nghi nhung khong duoc duyet)')

    doc.add_paragraph(
        "TAT CA cac loai nghi khac (phep nam, thai san, om AR2, cong tac, v.v.) "
        "KHONG tinh la vang khong phep."
    )

    add_example_box(doc, "Vi du 1 — Anh Tung (DAT):", [
        "Thang 3/2026: Anh Tung vang khong phep 1 ngay.",
        "1 ngay <= 2 ngay => DAT dieu kien C2."
    ])

    add_example_box(doc, "Vi du 2 — Chi Hoa (KHONG DAT):", [
        "Thang 3/2026: Chi Hoa vang khong phep 3 ngay.",
        "3 ngay > 2 ngay => KHONG DAT dieu kien C2.",
        "",
        "Luu y: Chi Hoa con nghi phep nam 2 ngay nua, nhung phep nam khong tinh la AR1.",
        "Chi co 3 ngay AR1 moi duoc dem."
    ])

    doc.add_page_break()

    # ---- C3 ----
    doc.add_heading("DIEU KIEN 3 (C3): NGAY LAM THUC TE", level=2)
    bold_para(doc, "Nguong: ", "> 0 ngay")
    bold_para(doc, "Ap dung cho: ", "TAT CA chuc vu Type-1 va Type-2")

    doc.add_paragraph(
        "Dieu kien nay rat don gian: ban phai co it nhat 1 ngay di lam thuc te trong thang. "
        "Neu ban nghi ca thang (vi du nghi thai san dai ngay, nghi om dai han), "
        "ban se khong nhan thuong thang do."
    )

    add_example_box(doc, "Vi du:", [
        "Chi Nga nghi thai san ca thang, ngay lam thuc te = 0 ngay.",
        "=> KHONG DAT dieu kien C3.",
        "",
        "Anh Binh di lam 1 ngay roi nghi om ca thang con lai.",
        "Ngay lam thuc te = 1 ngay > 0 => DAT dieu kien C3."
    ])

    # ---- C4 ----
    doc.add_heading("DIEU KIEN 4 (C4): NGAY LAM TOI THIEU", level=2)
    bold_para(doc, "Nguong: ", ">= 12 ngay (mac dinh)")
    bold_para(doc, "Ap dung cho: ", "TAT CA chuc vu Type-1 va Type-2")

    doc.add_paragraph(
        "Ban phai lam it nhat 12 ngay trong thang. Dieu kien nay de dam bao ban co mat du de duoc danh gia."
    )

    bold_para(doc, "Luu y dac biet: ",
              "Dieu kien nay CHI ap dung sau ngay 20 cua thang hien tai. "
              "Bao cao giua ky (truoc ngay 20) tu dong mien dieu kien nay. "
              "Ly do: Truoc ngay 20, nhieu nguoi chua lam du 12 ngay la binh thuong.")

    add_example_box(doc, "Vi du:", [
        "Anh Phuc lam duoc 10 ngay trong thang (vi moi vao lam ngay 15).",
        "Neu bao cao chay truoc ngay 20: C4 duoc MIEN (khong xet).",
        "Neu bao cao chay sau ngay 20: 10 ngay < 12 ngay => KHONG DAT C4."
    ])

    doc.add_page_break()

    # ---- C5 ----
    doc.add_heading("DIEU KIEN 5 (C5): LOI AQL CA NHAN", level=2)
    bold_para(doc, "Nguong: ", "= 0 loi (KHONG duoc co bat ky loi nao)")
    bold_para(doc, "Ap dung cho: ", "Assembly Inspector, RQC Assembly Inspector, AQL Inspector")

    doc.add_paragraph(
        "Day la dieu kien nghiem khac nhat: ban KHONG duoc co bat ky loi AQL nao trong thang. "
        "Chi 1 loi thoi la ban se khong nhan thuong."
    )

    doc.add_paragraph(
        "AQL la gi? AQL (Muc Chat luong Chap nhan duoc) la tieu chuan kiem tra quoc te. "
        "Khi kiem tra hang, neu san pham bi loi vuot qua muc cho phep, don hang bi tu choi (reject). "
        "Neu ban la nguoi kiem tra va de lot san pham loi, do la 'loi AQL ca nhan' cua ban."
    )

    add_example_box(doc, "Vi du 1 — Chi Yen (DAT):", [
        "Thang 3/2026: Chi Yen kiem tra nhieu don hang, khong co don nao bi tu choi vi loi cua chi.",
        "Loi AQL ca nhan = 0 => DAT dieu kien C5."
    ])

    add_example_box(doc, "Vi du 2 — Anh Dat (KHONG DAT):", [
        "Thang 3/2026: Anh Dat kiem tra don hang va 1 don bi tu choi vi anh de lot loi.",
        "Loi AQL ca nhan = 1 => KHONG DAT dieu kien C5.",
        "",
        "Luu y: Du chi co 1 loi, anh Dat van khong nhan thuong.",
        "Nhung neu loi la do doi san xuat (khong phai loi cua anh),",
        "quan ly co the xin Allowance (ngoai le) cho anh."
    ])

    # ---- C6 ----
    doc.add_heading("DIEU KIEN 6 (C6): LOI AQL LIEN TUC CA NHAN", level=2)
    bold_para(doc, "Nguong: ", "Khong duoc co loi AQL nhieu thang lien tiep")
    bold_para(doc, "Ap dung cho: ", "Assembly Inspector, RQC Assembly Inspector, AQL Inspector")

    doc.add_paragraph(
        "He thong kiem tra lich su AQL cua ban trong 3 thang gan nhat. "
        "Neu ban co loi AQL nhieu thang lien tiep, ban se khong nhan thuong."
    )

    bold_para(doc, "Quy tac theo nam:")
    bullet(doc, "Nam 2025: Chi 3 thang lien tuc co loi tro len => KHONG DAT")
    bullet(doc, "Tu nam 2026: Chi 2 thang lien tuc co loi tro len => KHONG DAT")

    add_example_box(doc, "Vi du (nam 2026):", [
        "Anh Cuong co loi AQL vao thang 1/2026 va thang 2/2026 (2 thang lien tiep).",
        "=> Thang 3/2026: KHONG DAT dieu kien C6.",
        "",
        "Chi Lan co loi AQL thang 1/2026, thang 2 khong loi, thang 3 co loi.",
        "=> Thang 3/2026: DAT dieu kien C6 (vi khong phai LIEN TIEP — thang 2 khong co loi)."
    ])

    doc.add_paragraph(
        "Bai hoc: Neu thang truoc ban co loi AQL, thang nay ban can lam that tot de cat chuoi lien tiep!"
    )

    doc.add_page_break()

    # ---- C7 ----
    doc.add_heading("DIEU KIEN 7 (C7): LOI AQL LIEN TUC NHOM/KHU VUC", level=2)
    bold_para(doc, "Nguong: ", "Cung quy tac nhu C6 nhung xet nhan vien duoi quyen")
    bold_para(doc, "Ap dung cho: ", "Line Leader, Auditor & Trainer")

    doc.add_paragraph(
        "Dieu kien nay khong xet ban truc tiep, ma xet nhung nhan vien ma ban phu trach:"
    )

    bullet(doc, "LINE LEADER: He thong kiem tra cac nhan vien duoi quyen cua ban. "
           "Neu bat ky nhan vien nao co loi AQL lien tuc (2+ thang lien tiep tu 2026), "
           "Line Leader se khong dat dieu kien nay.")

    bullet(doc, "AUDITOR & TRAINER: He thong kiem tra cac nhan vien trong khu vuc ban phu trach. "
           "Neu bat ky nhan vien nao co loi AQL lien tuc, Trainer se khong dat dieu kien nay.")

    add_example_box(doc, "Vi du:", [
        "Line Leader Chi Phuong co 8 nhan vien duoi quyen.",
        "Nhan vien A co loi AQL thang 2 va thang 3 (2 thang lien tiep).",
        "=> Chi Phuong KHONG DAT C7, du ban than chi khong co loi AQL nao.",
        "",
        "Bai hoc: Line Leader can theo doi va huong dan nhan vien de tranh loi lien tuc."
    ])

    bold_para(doc, "Luu y: ", "RQC Assembly Inspector (A1B) MIEN dieu kien nay tu thang 2/2026.")

    # ---- C8 ----
    doc.add_heading("DIEU KIEN 8 (C8): TY LE TU CHOI KHU VUC", level=2)
    bold_para(doc, "Nguong: ", "< 3.0% (mac dinh)")
    bold_para(doc, "Ap dung cho: ", "Auditor & Trainer, Model Master")

    doc.add_paragraph(
        "Dieu kien nay xet ty le tu choi AQL cua toan bo khu vuc ban phu trach:"
    )

    bullet(doc, "AUDITOR & TRAINER: Ty le tu choi cua toa nha ban phu trach phai < 3%.")
    bullet(doc, "MODEL MASTER: Ty le tu choi cua TOAN BO nha may phai < 3%.")

    add_example_box(doc, "Vi du — Trainer:", [
        "Chi Hien phu trach toa nha A2.",
        "Thang nay toa nha A2 kiem tra 200 don, 5 don bi tu choi.",
        "Ty le tu choi = 5/200 x 100 = 2.5% < 3% => DAT dieu kien C8.",
        "",
        "Thang sau toa nha A2 kiem tra 150 don, 6 don bi tu choi.",
        "Ty le tu choi = 6/150 x 100 = 4.0% > 3% => KHONG DAT dieu kien C8."
    ])

    add_example_box(doc, "Vi du — Model Master:", [
        "Anh Nam la Model Master.",
        "Thang nay toan bo nha may HWK kiem tra 1000 don, 25 don bi tu choi.",
        "Ty le tu choi = 25/1000 x 100 = 2.5% < 3% => DAT dieu kien C8."
    ])

    doc.add_page_break()

    # ---- C9 ----
    doc.add_heading("DIEU KIEN 9 (C9): TY LE DAT 5PRS", level=2)
    bold_para(doc, "Nguong: ", ">= 95% (mac dinh)")
    bold_para(doc, "Ap dung cho: ", "Assembly Inspector, RQC Assembly Inspector (A1B)")

    doc.add_paragraph(
        "5PRS la gi? 5PRS (Lay mau ngau nhien 5 doi) la cach kiem tra chat luong noi bo. "
        "Nguoi kiem tra lay ngau nhien 5 doi giay tu day chuyen de kiem tra. "
        "Neu 5 doi deu dat chuan -> Dat. Neu co doi bi loi -> Khong dat."
    )

    bold_para(doc, "Cong thuc: ", "Ty le dat = (So lan dat / Tong so lan kiem tra) x 100%")

    add_example_box(doc, "Vi du 1 — Chi Thu (DAT):", [
        "Thang 3/2026: Chi Thu kiem tra 5PRS 120 lan, trong do 118 lan dat.",
        "Ty le dat = (118/120) x 100 = 98.3%",
        "98.3% >= 95% => DAT dieu kien C9!"
    ])

    add_example_box(doc, "Vi du 2 — Anh Hieu (KHONG DAT):", [
        "Thang 3/2026: Anh Hieu kiem tra 5PRS 100 lan, trong do 90 lan dat.",
        "Ty le dat = (90/100) x 100 = 90.0%",
        "90.0% < 95% => KHONG DAT dieu kien C9!",
        "",
        "Anh Hieu can dat it nhat 95 lan tren 100 lan de dat C9."
    ])

    # ---- C10 ----
    doc.add_heading("DIEU KIEN 10 (C10): SO LUONG KIEM TRA 5PRS", level=2)
    bold_para(doc, "Nguong: ", ">= 100 doi (mac dinh)")
    bold_para(doc, "Ap dung cho: ", "Assembly Inspector (KHONG ap dung cho RQC Assembly Inspector A1B)")

    doc.add_paragraph(
        "Ban phai kiem tra it nhat 100 doi giay theo phuong phap 5PRS trong thang. "
        "Dieu kien nay dam bao ban thuc su lam viec kiem tra du luong."
    )

    add_example_box(doc, "Vi du 1 — Chi My (DAT):", [
        "Thang 3/2026: Chi My kiem tra 150 doi 5PRS.",
        "150 doi >= 100 doi => DAT dieu kien C10."
    ])

    add_example_box(doc, "Vi du 2 — Anh Long (KHONG DAT):", [
        "Thang 3/2026: Anh Long chi kiem tra 80 doi 5PRS (vi nghi nhieu ngay).",
        "80 doi < 100 doi => KHONG DAT dieu kien C10."
    ])

    bold_para(doc, "Luu y: ",
              "RQC Assembly Inspector (ma A1B) duoc MIEN dieu kien nay tu thang 2/2026. "
              "Ly do: RQC Inspector kiem tra quy trinh, khong phai kiem tra day chuyen 5PRS.")

    doc.add_page_break()

    # ---- APPENDIX 1 SUMMARY ----
    doc.add_heading("TOM TAT PHU LUC 1", level=2)

    doc.add_paragraph(
        "Duoi day la tom tat nhanh 10 dieu kien de ban de nho:"
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Dieu kien", "Yeu cau", "Diem then chot", "Meo dat"])

    summary_rows = [
        ["C1: Ty le di lam", ">= 88%", "Chi AR1 moi tinh la vang", "Di lam deu dan, xin phep dung quy dinh"],
        ["C2: Vang khong phep", "<= 2 ngay", "Chi dem ngay AR1", "Han che vang khong phep"],
        ["C3: Ngay lam thuc te", "> 0 ngay", "It nhat 1 ngay di lam", "Don gian — di lam it nhat 1 ngay"],
        ["C4: Ngay lam toi thieu", ">= 12 ngay", "Chi xet sau ngay 20", "Lam du ngay trong thang"],
        ["C5: Loi AQL ca nhan", "= 0 loi", "1 loi = truot", "Kiem tra ky luong, khong bo sot"],
        ["C6: AQL lien tuc ca nhan", "Khong lien tuc", "2+ thang lien tiep (2026+)", "Cat chuoi — lam tot thang nay"],
        ["C7: AQL lien tuc nhom", "Nhan vien khong loi", "Xet nhan vien duoi quyen", "Quan ly va huong dan nhan vien"],
        ["C8: Ty le tu choi KV", "< 3%", "Xet khu vuc phu trach", "Nang cao chat luong toan khu vuc"],
        ["C9: Ty le dat 5PRS", ">= 95%", "Kiem tra 5 doi ngau nhien", "Kiem tra can than moi doi"],
        ["C10: SL kiem tra 5PRS", ">= 100 doi", "A1B duoc mien", "Kiem tra du so luong moi ngay"],
    ]
    for row_data in summary_rows:
        add_row(tbl, row_data)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # PHU LUC 2
    # ══════════════════════════════════════════
    doc.add_heading("PHU LUC 2: MA TRAN AP DUNG DIEU KIEN THEO CHUC VU", level=1)

    doc.add_heading("Chuc vu Type-1", level=2)

    tbl = doc.add_table(rows=1, cols=11)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Chuc vu", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9", "C10"])

    full_matrix = [
        ["ASSEMBLY INSPECTOR", "Co", "Co", "Co", "Co", "Co", "Co", "—", "—", "Co", "Co"],
        ["RQC ASSEMBLY INSPECTOR\n(A1B, tu 02/2026)", "Co", "Co", "Co", "Co", "Co", "Co", "—", "—", "Co", "—"],
        ["AQL INSPECTOR", "Co", "Co", "Co", "Co", "Co", "—", "—", "—", "—", "—"],
        ["LINE LEADER", "Co", "Co", "Co", "Co", "—", "—", "Co", "—", "—", "—"],
        ["AUDITOR & TRAINER", "Co", "Co", "Co", "Co", "—", "—", "Co", "Co", "—", "—"],
        ["MODEL MASTER", "Co", "Co", "Co", "Co", "—", "—", "—", "Co", "—", "—"],
        ["Quan ly (GL, Supervisor,\nManager)", "Co", "Co", "Co", "Co", "—", "—", "—", "—", "—", "—"],
    ]
    for m in full_matrix:
        add_row(tbl, m)

    doc.add_paragraph()
    doc.add_paragraph(
        '"Co" = dieu kien ap dung cho chuc vu nay, phai dat. '
        '"—" = khong ap dung, duoc mien.'
    )

    doc.add_heading("Chuc vu Type-2", level=2)
    doc.add_paragraph(
        "Tat ca Type-2 chi ap dung 4 dieu kien chuyen can (C1-C4). "
        "So tien thuong theo Phan 5.6 va 5.7."
    )

    doc.add_heading("Chuc vu Type-3", level=2)
    doc.add_paragraph(
        "Type-3 khong co dieu kien ap dung va khong nhan thuong. "
        "Tu dong chuyen Type-2 sau khoang 3 thang."
    )

    doc.add_page_break()

    # PHU LUC 3
    doc.add_heading("PHU LUC 3: TRANG THAI TALENT POOL QIP HWK", level=1)
    doc.add_paragraph(
        "Talent Pool QIP HWK la nhom nhan vien xuat sac duoc quan ly chon. "
        "Thanh vien Talent Pool nhan thuong bo sung hang thang."
    )
    bullet(doc, "Tu dong ap dung: Co — Thuong tu dong cung thuong thuong xuyen")
    bullet(doc, "Cong don: Co — Cong them vao thuong Chuong trinh 1")
    bullet(doc, "Thoi gian thanh toan: Cung ky thuong hang thang")
    bullet(doc, "Yeu cau dieu kien: Khong — nhan thuong bat ke dat/khong dat dieu kien")
    bullet(doc, "Loai tru: Nhan vien da nghi tu dong bi loai; yeu cau trang thai hoat dong")

    add_example_box(doc, "Vi du:", [
        "Chi Hieu duoc chon vao Talent Pool voi muc thuong 500,000 VND/thang.",
        "Thang nay chi cung dat du dieu kien va nhan thuong Chuong trinh 1 = 350,000 VND.",
        "=> Tong thuong thang nay = 350,000 + 500,000 = 850,000 VND.",
        "",
        "Thang sau chi khong dat dieu kien nen thuong Chuong trinh 1 = 0 VND.",
        "Nhung thuong Talent Pool van duoc nhan = 500,000 VND.",
        "=> Tong thuong = 0 + 500,000 = 500,000 VND."
    ])

    doc.add_page_break()

    # PHU LUC 4
    doc.add_heading("PHU LUC 4: BAN DO R&R DOI QIP", level=1)
    doc.add_paragraph("Vui long tham khao tai lieu Ban do R&R Doi QIP rieng cho danh sach 69 cong viec.")

    # PHU LUC 5
    doc.add_heading("PHU LUC 5: BO SUU TAP DU LIEU FIRESTORE", level=1)
    doc.add_paragraph("He thong V10 su dung cac bo suu tap Firestore sau:")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    make_header_row(tbl, ["Bo suu tap", "Muc dich", "Quyen truy cap"])
    collections = [
        ["employees/{monthYear}/all_data/data", "Du lieu nhan vien + thuong", "Xac thuc: doc, Admin: ghi"],
        ["dashboard_summary/{monthYear}", "Thong ke tom tat KPI", "Xac thuc: doc, Admin: ghi"],
        ["thresholds/{monthYear}", "7 nguong cau hinh", "Xac thuc: doc, Admin: ghi"],
        ["threshold_history/{autoId}", "Nhat ky kiem tra khong sua doi", "Xac thuc: doc, Admin: tao"],
        ["allowances/{monthYear}/items/{docId}", "Ngoai le allowance", "Xac thuc: doc, Admin: ghi"],
        ["system/config", "Cai dat he thong", "Xac thuc: doc, Admin: ghi"],
        ["configs/{document}", "Anh xa chuc vu, cau hinh", "Xac thuc: doc, Admin: ghi"],
        ["system_feedback/{autoId}", "Phan hoi / theo doi van de", "Xac thuc: doc/tao, Admin: day du"],
        ["pendingNotifications/{autoId}", "Hang doi thong bao email", "Xac thuc: tao, Admin: doc"],
        ["config/email", "Thong tin SMTP", "Chi Admin"],
        ["email_logs/{autoId}", "Nhat ky gui email", "Chi Admin"],
    ]
    for c in collections:
        add_row(tbl, c)

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Het —")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Neu ban co bat ky thac mac nao ve chinh sach thuong, "
        "hay lien he quan ly QIP hoac gui phan hoi qua he thong."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(89, 89, 89)

    return doc


if __name__ == "__main__":
    doc = create_policy_vi()
    output_path = os.path.expanduser(
        "~/Downloads/2025 QIP INCENTIVE POLICY Version3_Mar2026-Feb2027_Vi.docx"
    )
    doc.save(output_path)
    print(f"Policy document (Vietnamese) saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path):,} bytes")
